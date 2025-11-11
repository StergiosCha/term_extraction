"""
Complete TextCraft AI Main with Multi-LLM Support + Terminology Scraper
Includes the eleto.gr scraper and RAG system
"""

import os
from dotenv import load_dotenv
load_dotenv()

import logging
import asyncio
import json
import re
import unicodedata
import sqlite3
import google.generativeai as genai
import openai
import anthropic
from datetime import datetime, timedelta, date
from enum import Enum
from typing import List, Dict, Optional, Tuple, Any
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from dataclasses import dataclass
import pickle
import PyPDF2
from io import BytesIO, StringIO
import docx
import csv
import aiohttp
import fitz # PyMuPDF
import chardet
from pdfminer.high_level import extract_text_to_fp
from tqdm import tqdm

# FastAPI imports
from fastapi import FastAPI, Request, Form, File, UploadFile, HTTPException, Depends, status
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles

# SQLAlchemy imports
from sqlalchemy import create_engine, Column, Integer, String, Boolean, DateTime, event
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session as SQLASessionType
from sqlalchemy.sql import func

# Other utility imports
from pathlib import Path
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Import your local modules
from models import User, Session as DBSessionModel, Subscription, create_tables, get_db, check_daily_limits, get_remaining_corrections, update_user_tier_features, migrate_privacy_columns, CustomGlossary, TranslationHistory, FavoriteTranslation, BatchTranslation
from auth import AuthService
from email_service import EmailService
from typing import List, Dict
import stripe

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- FIX: Define absolute paths for file serving ---
# Get the absolute path to the directory where this script is located
BASE_DIR = Path(__file__).parent
STATIC_DIR = BASE_DIR / "static"
TEMPLATES_DIR = BASE_DIR / "templates"

# Initialize FastAPI app
app = FastAPI(title="TextCraft AI - Professional Text Editor")

# --- FIX: Mount static files using the absolute path ---
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# LLM Provider Support - Expanded Model List
class LLMProvider(str, Enum):
    # Google Gemini Models
    GEMINI_2_5_PRO = "gemini-2.5-pro"
    GEMINI_2_0_FLASH_EXP = "gemini-2.0-flash-exp"
    GEMINI_1_5_PRO = "gemini-1.5-pro"
    GEMINI_1_5_FLASH = "gemini-1.5-flash"
    GEMINI_1_5_FLASH_8B = "gemini-1.5-flash-8b"
    
    # Anthropic Claude Models
    CLAUDE_3_7_SONNET = "claude-3-7-sonnet-20250219"
    CLAUDE_3_5_SONNET = "claude-3-5-sonnet-20241022"
    CLAUDE_3_5_HAIKU = "claude-3-5-haiku-20241022"
    CLAUDE_3_OPUS = "claude-3-opus-20240229"
    CLAUDE_3_SONNET = "claude-3-sonnet-20240229"
    CLAUDE_3_HAIKU = "claude-3-haiku-20240307"
    
    # OpenAI GPT Models
    GPT_4O = "gpt-4o"
    GPT_4O_MINI = "gpt-4o-mini"
    GPT_4_TURBO = "gpt-4-turbo"
    GPT_4 = "gpt-4"
    GPT_3_5_TURBO = "gpt-3.5-turbo"
    O1_PREVIEW = "o1-preview"
    O1_MINI = "o1-mini"

# Model Configuration
MODEL_CONFIG = {
    # Gemini Models
    LLMProvider.GEMINI_2_5_PRO: {
        "name": "Gemini 2.5 Pro",
        "provider": "gemini",
        "model_id": "gemini-2.0-flash-exp",  # API model ID
        "description": "Google's latest advanced multimodal model",
        "max_tokens": 8192,
        "supports_streaming": True
    },
    LLMProvider.GEMINI_2_0_FLASH_EXP: {
        "name": "Gemini 2.0 Flash (Experimental)",
        "provider": "gemini",
        "model_id": "gemini-2.0-flash-exp",
        "description": "Fast experimental multimodal model",
        "max_tokens": 8192,
        "supports_streaming": True
    },
    LLMProvider.GEMINI_1_5_PRO: {
        "name": "Gemini 1.5 Pro",
        "provider": "gemini",
        "model_id": "gemini-1.5-pro",
        "description": "Advanced reasoning with long context",
        "max_tokens": 8192,
        "supports_streaming": True
    },
    LLMProvider.GEMINI_1_5_FLASH: {
        "name": "Gemini 1.5 Flash",
        "provider": "gemini",
        "model_id": "gemini-1.5-flash",
        "description": "Fast and efficient model",
        "max_tokens": 8192,
        "supports_streaming": True
    },
    LLMProvider.GEMINI_1_5_FLASH_8B: {
        "name": "Gemini 1.5 Flash 8B",
        "provider": "gemini",
        "model_id": "gemini-1.5-flash-8b",
        "description": "Lightweight 8B parameter model",
        "max_tokens": 8192,
        "supports_streaming": True
    },
    
    # Claude Models
    LLMProvider.CLAUDE_3_7_SONNET: {
        "name": "Claude 3.7 Sonnet",
        "provider": "anthropic",
        "model_id": "claude-3-7-sonnet-20250219",
        "description": "Anthropic's latest advanced reasoning model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.CLAUDE_3_5_SONNET: {
        "name": "Claude 3.5 Sonnet",
        "provider": "anthropic",
        "model_id": "claude-3-5-sonnet-20241022",
        "description": "Advanced reasoning and coding",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.CLAUDE_3_5_HAIKU: {
        "name": "Claude 3.5 Haiku",
        "provider": "anthropic",
        "model_id": "claude-3-5-haiku-20241022",
        "description": "Fast and efficient Claude model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.CLAUDE_3_OPUS: {
        "name": "Claude 3 Opus",
        "provider": "anthropic",
        "model_id": "claude-3-opus-20240229",
        "description": "Most capable Claude model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.CLAUDE_3_SONNET: {
        "name": "Claude 3 Sonnet",
        "provider": "anthropic",
        "model_id": "claude-3-sonnet-20240229",
        "description": "Balanced performance and speed",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.CLAUDE_3_HAIKU: {
        "name": "Claude 3 Haiku",
        "provider": "anthropic",
        "model_id": "claude-3-haiku-20240307",
        "description": "Fastest Claude model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    
    # OpenAI Models
    LLMProvider.GPT_4O: {
        "name": "GPT-4o",
        "provider": "openai",
        "model_id": "gpt-4o",
        "description": "OpenAI's latest multimodal model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.GPT_4O_MINI: {
        "name": "GPT-4o Mini",
        "provider": "openai",
        "model_id": "gpt-4o-mini",
        "description": "Faster and cheaper GPT-4o variant",
        "max_tokens": 16384,
        "supports_streaming": True
    },
    LLMProvider.GPT_4_TURBO: {
        "name": "GPT-4 Turbo",
        "provider": "openai",
        "model_id": "gpt-4-turbo",
        "description": "Enhanced GPT-4 with longer context",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.GPT_4: {
        "name": "GPT-4",
        "provider": "openai",
        "model_id": "gpt-4",
        "description": "Original GPT-4 model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.GPT_3_5_TURBO: {
        "name": "GPT-3.5 Turbo",
        "provider": "openai",
        "model_id": "gpt-3.5-turbo",
        "description": "Fast and cost-effective model",
        "max_tokens": 4096,
        "supports_streaming": True
    },
    LLMProvider.O1_PREVIEW: {
        "name": "O1 Preview",
        "provider": "openai",
        "model_id": "o1-preview",
        "description": "OpenAI's reasoning model",
        "max_tokens": 4096,
        "supports_streaming": False
    },
    LLMProvider.O1_MINI: {
        "name": "O1 Mini",
        "provider": "openai",
        "model_id": "o1-mini",
        "description": "Smaller reasoning model",
        "max_tokens": 4096,
        "supports_streaming": False
    },
}

class MultiLLMManager:
    """Manages multiple LLM providers - USERS MUST PROVIDE THEIR OWN API KEYS"""
    
    def __init__(self):
        self.default_provider = LLMProvider.GEMINI_2_5_PRO
        # Store user API keys per session (in-memory for now, can be moved to DB)
        # NO SYSTEM API KEYS - users must provide their own
        self.user_api_keys = {}  # {session_id: {provider_type: api_key}}
    
    def set_user_api_key(self, session_id: str, provider_type: str, api_key: str):
        """Set user API key for a session"""
        if session_id not in self.user_api_keys:
            self.user_api_keys[session_id] = {}
        self.user_api_keys[session_id][provider_type] = api_key
        logger.info(f"Set user API key for {provider_type} (session: {session_id[:8]}...)")
    
    def get_api_key(self, session_id: Optional[str], provider_type: str) -> Optional[str]:
        """Get API key - ONLY user-provided keys, NO system fallback"""
        if session_id and session_id in self.user_api_keys:
            return self.user_api_keys[session_id].get(provider_type)
        return None
    
    def get_available_models(self, session_id: Optional[str] = None) -> Dict[str, Dict[str, Any]]:
        """Get list of all available models with their availability status"""
        available = {}
        
        for model_enum, config in MODEL_CONFIG.items():
            provider_type = config["provider"]
            api_key = self.get_api_key(session_id, provider_type)
            
            available[model_enum.value] = {
                "name": config["name"],
                "description": config["description"],
                "provider": provider_type,
                "model_id": config["model_id"],
                "available": api_key is not None,
                "has_user_key": session_id and session_id in self.user_api_keys and provider_type in self.user_api_keys.get(session_id, {}),
                "max_tokens": config["max_tokens"],
                "supports_streaming": config["supports_streaming"]
            }
        
        return available
    
    def is_model_available(self, model: str, session_id: Optional[str] = None) -> bool:
        """Check if a specific model is available (has API key)"""
        try:
            model_enum = LLMProvider(model)
            config = MODEL_CONFIG[model_enum]
            provider_type = config["provider"]
            api_key = self.get_api_key(session_id, provider_type)
            return api_key is not None
        except (ValueError, KeyError):
            return False
    
    async def generate_with_provider(self, prompt: str, provider: str, timeout: int = 30, session_id: Optional[str] = None) -> str:
        """Generate text using specified provider/model"""
        
        try:
            model_enum = LLMProvider(provider)
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unknown model: '{provider}'"
            )
        
        config = MODEL_CONFIG.get(model_enum)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Model configuration not found: '{provider}'"
            )
        
        provider_type = config["provider"]
        model_id = config["model_id"]
        api_key = self.get_api_key(session_id, provider_type)
        
        if not api_key:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"API key required for {provider_type}. Please provide your own API key in the settings (click the ⚙️ icon)."
            )
        
        try:
            if provider_type == "gemini":
                return await self._generate_gemini(prompt, api_key, model_id, timeout)
            elif provider_type == "anthropic":
                return await self._generate_claude(prompt, api_key, model_id, timeout, config["max_tokens"])
            elif provider_type == "openai":
                return await self._generate_openai(prompt, api_key, model_id, timeout, config["max_tokens"])
            else:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Unknown provider type: {provider_type}"
                )
                
        except asyncio.TimeoutError:
            logger.error(f"LLM generation timed out after {timeout} seconds with {provider}")
            raise HTTPException(
                status_code=status.HTTP_408_REQUEST_TIMEOUT,
                detail="Request timed out. Please try a shorter text or different model."
            )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"LLM generation error with {provider}: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Generation failed with {provider}: {str(e)}"
            )
    
    async def _generate_gemini(self, prompt: str, api_key: str, model_id: str, timeout: int) -> str:
        """Generate with Gemini"""
        loop = asyncio.get_event_loop()
        
        def call_gemini():
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_id)
            response = model.generate_content(prompt)
            return response.text
        
        response_task = loop.run_in_executor(None, call_gemini)
        response = await asyncio.wait_for(response_task, timeout=timeout)
        return response
    
    async def _generate_claude(self, prompt: str, api_key: str, model_id: str, timeout: int, max_tokens: int) -> str:
        """Generate with Claude"""
        loop = asyncio.get_event_loop()
        
        def call_claude():
            client = anthropic.Anthropic(api_key=api_key)
            message = client.messages.create(
                model=model_id,
                max_tokens=max_tokens,
                messages=[{"role": "user", "content": prompt}]
            )
            return message.content[0].text
        
        response_task = loop.run_in_executor(None, call_claude)
        response = await asyncio.wait_for(response_task, timeout=timeout)
        return response
    
    async def _generate_openai(self, prompt: str, api_key: str, model_id: str, timeout: int, max_tokens: int) -> str:
        """Generate with OpenAI"""
        loop = asyncio.get_event_loop()
        
        def call_openai():
            client = openai.OpenAI(api_key=api_key)
            # O1 models use different API
            if model_id.startswith("o1"):
                response = client.chat.completions.create(
                    model=model_id,
                    messages=[{"role": "user", "content": prompt}]
                )
            else:
                response = client.chat.completions.create(
                    model=model_id,
                    messages=[{"role": "user", "content": prompt}],
                    max_completion_tokens=max_tokens,
                    temperature=1
                )
            return response.choices[0].message.content
        
        response_task = loop.run_in_executor(None, call_openai)
        response = await asyncio.wait_for(response_task, timeout=timeout)
        return response

# Initialize global LLM manager
llm_manager = MultiLLMManager()

# Enhanced generation functions
async def generate_with_timeout_multi(prompt: str, provider: str = None, timeout: int = 30, session_id: Optional[str] = None) -> str:
    """Generate with timeout using specified provider"""
    if provider is None:
        provider = llm_manager.default_provider.value
    
    return await llm_manager.generate_with_provider(prompt, provider, timeout, session_id)

async def generate_with_timeout(prompt: str, timeout: int = 30, session_id: Optional[str] = None) -> str:
    """Original function - uses default provider (Gemini)"""
    return await generate_with_timeout_multi(prompt, llm_manager.default_provider.value, timeout, session_id)

# TERMINOLOGY SCRAPER AND RAG SYSTEM

@dataclass
class TerminologyEntry:
    """Represents a terminology entry"""
    term_greek: str
    term_english: str
    definition_greek: str
    definition_english: str
    category: str
    source_url: str
    context: str = ""
    domain: str = ""
    confidence: float = 1.0

class EletoDocumentScraper:
    """Aggressive scraper for all eleto.gr documents"""
    
    def __init__(self):
        self.base_url = "https://eleto.gr"
        self.session = None
        self.scraped_urls = set()
        self.document_cache = {}
        
    async def init_session(self):
        if not self.session or self.session.closed:
            connector = aiohttp.TCPConnector(
                limit=5,
                limit_per_host=2,
                ttl_dns_cache=300,
                use_dns_cache=True,
            )
            
            timeout = aiohttp.ClientTimeout(
                total=60,
                connect=30,
                sock_read=30
            )
            
            self.session = aiohttp.ClientSession(
                connector=connector,
                timeout=timeout,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.9,el;q=0.8',
                    'Accept-Encoding': 'gzip, deflate',  # Removed 'br' to avoid brotli
                    'DNT': '1',
                    'Connection': 'keep-alive',
                    'Upgrade-Insecure-Requests': '1',
                    'Cache-Control': 'max-age=0'
                }
            )
            
    async def close_session(self):
        if self.session and not self.session.closed:
            await self.session.close()
            
    async def fetch_page(self, url: str, retries: int = 3) -> Optional[str]:
        await self.init_session()
        
        for attempt in range(retries):
            try:
                await asyncio.sleep(2 + attempt)
                
                async with self.session.get(url) as response:
                    if response.status == 200:
                        content = await response.text()
                        logger.info(f"Fetched: {url} (attempt {attempt + 1})")
                        return content
                    elif response.status == 403:
                        logger.warning(f"Access denied: {url}")
                        return None
                    elif response.status == 429:
                        logger.warning(f"Rate limited: {url}, waiting...")
                        await asyncio.sleep(10)
                        continue
                    else:
                        logger.warning(f"HTTP {response.status}: {url}")
                        
            except aiohttp.ClientError as e:
                logger.warning(f"Attempt {attempt + 1} failed for {url}: {e}")
                if attempt < retries - 1:
                    await asyncio.sleep(5 * (attempt + 1))
                else:
                    logger.error(f"All attempts failed for {url}")
                    
            except Exception as e:
                logger.error(f"Unexpected error for {url}: {e}")
                break
                
        return None
    async def download_file(self, url: str) -> Optional[bytes]:
        """Download PDF/DOC files"""
        await self.init_session()
        try:
            async with self.session.get(url) as response:
                if response.status == 200:
                    content = await response.read()
                    logger.info(f"📄 Downloaded file: {url}")
                    return content
                else:
                    logger.warning(f"❌ Failed to download {url}: {response.status}")
                    return None
        except Exception as e:
            logger.error(f"❌ Download error {url}: {e}")
            return None

    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF and a fallback to pdfminer.six."""
        text = ""
        try:
            # Attempt 1: Try PyMuPDF
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()

            # Simple sanity check to see if extraction was successful
            if len(text) > 100:
                logger.info("✅ PyMuPDF extraction successful.")
                return text

        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying fallback: {e}")
            
        try:
            # Attempt 2: Fallback to pdfminer.six
            output = BytesIO()
            extract_text_to_fp(BytesIO(pdf_content), output)
            text = output.getvalue().decode('utf-8')
            logger.info("✅ Fallback extraction with pdfminer.six succeeded.")
            return text
        except Exception as e:
            logger.error(f"❌ Fallback PDF extraction error: {e}")
            return ""

    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            docx_file = BytesIO(docx_content)
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    async def scrape_all_links_from_page(self, url: str) -> List[str]:
        """Get ALL links from a page"""
        html = await self.fetch_page(url)
        if not html:
            return []
            
        soup = BeautifulSoup(html, 'html.parser')
        links = set()
        
        for link in soup.find_all('a', href=True):
            href = link.get('href')
            if href:
                full_url = urljoin(url, href)
                if 'eleto.gr' in full_url:
                    links.add(full_url)
        
        return list(links)

    async def find_all_document_links(self, start_urls: List[str]) -> List[str]:
        """Recursively find ALL document links"""
        all_links = set(start_urls)
        to_process = list(start_urls)
        processed = set()
        
        while to_process:
            current_url = to_process.pop(0)
            if current_url in processed:
                continue
                
            processed.add(current_url)
            logger.info(f"🔍 Exploring: {current_url}")
            
            page_links = await self.scrape_all_links_from_page(current_url)
            
            for link in page_links:
                if link not in all_links:
                    all_links.add(link)
                    
                    if not self.is_document_file(link) and self.is_relevant_page(link):
                        to_process.append(link)
            
            await asyncio.sleep(1)
            
        return list(all_links)

    def is_document_file(self, url: str) -> bool:
        """Check if URL points to a document file"""
        url_lower = url.lower()
        return any(ext in url_lower for ext in ['.pdf', '.doc', '.docx', '.txt', '.rtf'])

    def is_relevant_page(self, url: str) -> bool:
        """Check if page is relevant for terminology"""
        url_lower = url.lower()
        relevant_keywords = [
            'lexika', 'glossaria', 'vaseis-oron', 'dictionary', 'glossary', 
            'terminology', 'terms', 'lexicon', 'vocabulary', 'orologikoi-poroi'
        ]
        return any(keyword in url_lower for keyword in relevant_keywords)

    async def scrape_everything(self) -> Dict[str, str]:
        """Scrape ALL content from eleto.gr terminology pages"""
        start_urls = [
            "https://eleto.gr/el/orologikoi-poroi/lexika-kai-glossaria/",
            "https://eleto.gr/el/orologikoi-poroi/vaseis-oron/"
        ]
        
        print(f"🚀 Starting comprehensive scraping from {len(start_urls)} URLs...")
        logger.info("🚀 Starting comprehensive scraping...")
        
        all_links = await self.find_all_document_links(start_urls)
        print(f"📋 Found {len(all_links)} total links")
        logger.info(f"📋 Found {len(all_links)} total links")
        
        document_files = [url for url in all_links if self.is_document_file(url)]
        html_pages = [url for url in all_links if not self.is_document_file(url)]
        
        print(f"📄 Document files: {len(document_files)}")
        print(f"🌐 HTML pages: {len(html_pages)}")
        logger.info(f"📄 Document files: {len(document_files)}")
        logger.info(f"🌐 HTML pages: {len(html_pages)}")
        
        scraped_content = {}
        
        for i, url in enumerate(html_pages):
            print(f"Processing HTML page {i+1}/{len(html_pages)}: {url}")
            content = await self.scrape_html_content(url)
            if content:
                scraped_content[url] = content
                print(f"✅ Scraped HTML content from {url}")
                
        for i, url in enumerate(document_files):
            print(f"Processing document {i+1}/{len(document_files)}: {url}")
            content = await self.scrape_document_file(url)
            if content:
                scraped_content[url] = content
                print(f"✅ Scraped document content from {url}")
                
        print(f"✅ Total scraped: {len(scraped_content)} documents")
        logger.info(f"✅ Scraped {len(scraped_content)} documents total")
        return scraped_content

    async def scrape_html_content(self, url: str) -> Optional[str]:
        """Extract terminology content from HTML page"""
        html = await self.fetch_page(url)
        if not html:
            return None
            
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup(['nav', 'footer', 'header', 'aside', 'script', 'style']):
            element.decompose()
            
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|article'))
        
        if main_content:
            text = main_content.get_text(separator='\n', strip=True)
        else:
            text = soup.get_text(separator='\n', strip=True)
            
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        cleaned_text = '\n'.join(lines)
        
        if len(cleaned_text) > 100:
            return cleaned_text
            
        return None

    async def scrape_document_file(self, url: str) -> Optional[str]:
        """Download and extract text from document files"""
        file_content = await self.download_file(url)
        if not file_content:
            return None
            
        url_lower = url.lower()
        
        if '.pdf' in url_lower:
            return self.extract_text_from_pdf(file_content)
        elif '.docx' in url_lower:
            return self.extract_text_from_docx(file_content)
        elif '.doc' in url_lower:
            try:
                # Assuming .doc files can be processed like .docx
                return self.extract_text_from_docx(file_content)
            except Exception:
                # Fallback decoding for .doc if docx library fails
                result = chardet.detect(file_content)
                encoding = result['encoding']
                if encoding:
                    return file_content.decode(encoding, errors='ignore')
                return ""
        elif any(ext in url_lower for ext in ['.txt', '.rtf']):
            # Use chardet for automatic encoding detection
            result = chardet.detect(file_content)
            encoding = result['encoding']
            if encoding:
                return file_content.decode(encoding, errors='ignore')
            return ""
            
        return None

class TerminologyRAGSystem:
    """RAG system for terminology-aware translation"""
    
    def __init__(self, db_path: str = "terminology_rag.db"):
        self.data_dir = Path("terminology_data")
        self.data_dir.mkdir(exist_ok=True)
        self.embedding_model = None
        self.index = None
        self.document_chunks = []
        self.chunk_sources = []
        
    def load_embedding_model(self):
        """Load multilingual embedding model"""
        try:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Embedding model loaded")
        except Exception as e:
            logger.error(f"❌ Failed to load embedding model: {e}")
            try:
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("✅ Fallback embedding model loaded")
            except Exception as e2:
                logger.error(f"❌ Failed to load fallback model: {e2}")

    def chunk_text_by_lines(self, text_content: str, source_url: str, min_lines: int = 5, max_lines: int = 20, overlap_lines: int = 3) -> List[str]:
        """
        Simple: add source URL and first 3 lines to each chunk for context
        """
        
        lines = text_content.split('\n')
        
        # Get first 3 meaningful lines as document context
        document_context_lines = []
        for line in lines[:10]:  # Check first 10 lines to find 3 meaningful ones
            line = line.strip()
            if line and len(line) > 5:  # Skip very short lines
                document_context_lines.append(line)
            if len(document_context_lines) >= 3:
                break
        
        # Create header with source and document context
        source_header = f"SOURCE: {source_url}"
        if document_context_lines:
            source_header += "\nDOCUMENT START:\n" + "\n".join(document_context_lines)
        source_header += "\n---"
        
        # Do the actual chunking logic
        clean_lines = [line.strip() for line in lines if line.strip()]
        
        if not clean_lines:
            return []

        if len(clean_lines) <= min_lines:
            joined_lines = '\n'.join(clean_lines)
            return [f"{source_header}\n{joined_lines}"]

        chunks = []
        i = 0
        safety_counter = 0
        max_iterations = len(clean_lines) * 2

        while i < len(clean_lines) and safety_counter < max_iterations:
            safety_counter += 1

            chunk_end = min(i + max_lines, len(clean_lines))

            # Ensure minimum lines
            if chunk_end - i < min_lines and chunk_end < len(clean_lines):
                chunk_end = min(i + min_lines, len(clean_lines))

            # Extract chunk
            chunk_lines = clean_lines[i:chunk_end]
            chunk_text = '\n'.join(chunk_lines).strip()

            if chunk_text and len(chunk_lines) >= min_lines:
                # Add source header to each chunk
                full_chunk = f"{source_header}\n{chunk_text}"
                chunks.append(full_chunk)

            # Move forward with overlap
            next_i = chunk_end - overlap_lines if overlap_lines > 0 else chunk_end
            if next_i <= i:
                next_i = i + 1
            i = next_i

        return chunks

    # Replace the existing chunking in load_and_process_saved_files
    def load_and_process_with_simple_metadata(self):
        """Just add source URLs to chunks, nothing fancy"""
        
        saved_files = list(self.data_dir.glob("*.txt"))
        print(f"Processing {len(saved_files)} files with source info...")
        
        if not self.embedding_model:
            self.load_embedding_model()
        
        all_chunks = []
        chunk_sources = []
        
        for file_path in saved_files:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                # Extract URL from first line if it exists
                lines = content.split('\n', 2)
                if lines[0].startswith('SOURCE:'):
                    url = lines[0].replace('SOURCE: ', '')
                    text_content = lines[2] if len(lines) > 2 else content
                else:
                    url = str(file_path)
                    text_content = content
                
                # Create chunks with source info
                chunks = self.create_chunk_with_source_info(text_content, url)
                all_chunks.extend(chunks)
                chunk_sources.extend([url] * len(chunks))
                
            except Exception as e:
                print(f"Failed to process {file_path}: {e}")
        
        self.document_chunks = all_chunks
        self.chunk_sources = chunk_sources
        
        print(f"Created {len(all_chunks)} chunks with source info")
        
        # Embed and index as usual
        batch_size = 16
        embeddings = []
        
        for i in range(0, len(self.document_chunks), batch_size):
            batch = self.document_chunks[i:i+batch_size]
            try:
                batch_embeddings = self.embedding_model.encode(batch, batch_size=batch_size, show_progress_bar=False)
                embeddings.extend(batch_embeddings)
            except Exception as e:
                print(f"Encoding error: {e}")
        
        if embeddings:
            self.build_faiss_index_from_embeddings(embeddings)
            print("Index built with source-aware chunks")
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
                
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        return chunks
    def save_index_and_metadata(self):
        """Save FAISS index and metadata to disk"""
        try:
            index_path = Path("terminology.index")
            chunks_path = Path("chunks.pkl")
            sources_path = Path("sources.pkl")

            faiss.write_index(self.index, str(index_path))

            with open(chunks_path, "wb") as f:
                pickle.dump(self.document_chunks, f)

            with open(sources_path, "wb") as f:
                pickle.dump(self.chunk_sources, f)

            logger.info(f"✅ Saved FAISS index and metadata")
        except Exception as e:
            logger.error(f"❌ Failed to save index: {e}")

    def load_existing_index(self) -> bool:
        """Load FAISS index and metadata from disk"""
        try:
            index_path = Path("terminology.index")
            chunks_path = Path("chunks.pkl")
            sources_path = Path("sources.pkl")

            if not index_path.exists() or not chunks_path.exists() or not sources_path.exists():
                logger.warning("⚠️ Saved index or metadata not found")
                return False

            self.index = faiss.read_index(str(index_path))

            with open(chunks_path, "rb") as f:
                self.document_chunks = pickle.load(f)

            with open(sources_path, "rb") as f:
                self.chunk_sources = pickle.load(f)

            logger.info(f"✅ Loaded FAISS index with {len(self.document_chunks)} chunks")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to load index: {e}")
            return False

    def save_files_only(self, scraped_content: Dict[str, str]):
        """Save all files to disk without any processing"""
        print(f"Saving {len(scraped_content)} files to disk...")
        
        for i, (url, content) in enumerate(scraped_content.items()):
            if len(content) < 50:
                continue
                
            # Create unique filename
            url_hash = abs(hash(url)) % 100000
            filename = f"doc_{i:04d}_{url_hash}"
            file_path = self.data_dir / f"{filename}.txt"
            
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(f"SOURCE: {url}\n\n{content}")
                
                if i % 100 == 0:
                    print(f"Saved {i}/{len(scraped_content)} files...")
                    
            except Exception as e:
                print(f"Failed to save {filename}: {e}")
        
        saved_files = list(self.data_dir.glob("*.txt"))
        print(f"Total files saved: {len(saved_files)}")
        return len(saved_files)


    def load_and_process_saved_files(self):
        """Load saved files, create chunks, encode them into embeddings, and build FAISS index."""
        
        # Check for cached embeddings first
        embeddings_cache_file = "embeddings_cache.pkl"
        chunks_cache_file = "chunks_cache.pkl"
        sources_cache_file = "sources_cache.pkl"
        
        if (os.path.exists(embeddings_cache_file) and 
            os.path.exists(chunks_cache_file) and 
            os.path.exists(sources_cache_file)):
            
            print("Loading cached data...")
            try:
                with open(embeddings_cache_file, 'rb') as f:
                    embeddings = pickle.load(f)
                with open(chunks_cache_file, 'rb') as f:
                    self.document_chunks = pickle.load(f)
                with open(sources_cache_file, 'rb') as f:
                    self.chunk_sources = pickle.load(f)
                    
                print(f"Loaded {len(embeddings)} cached embeddings and {len(self.document_chunks)} chunks")
                
                # Skip directly to FAISS index building
                self.build_faiss_index_from_embeddings(embeddings)
                return
                
            except Exception as e:
                print(f"Failed to load cache: {e}")
                print("Regenerating from scratch...")
        
        # Original processing code
        saved_files = list(self.data_dir.glob("*.txt"))
        print(f"Loading {len(saved_files)} saved files...")

        # Process ALL files without filtering
        if not self.embedding_model:
            self.load_embedding_model()

        all_chunks = []
        chunk_sources = []

        # Process all files with progress bar
        for file_path in tqdm(saved_files, desc="Processing files", unit="file"):
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()

                # Extract URL from first line
                lines = content.split('\n', 2)
                url = lines[0].replace('SOURCE: ', '') if lines[0].startswith('SOURCE:') else str(file_path)
                text_content = lines[2] if len(lines) > 2 else content

                chunks = self.chunk_text_by_lines(text_content, url, min_lines=5, max_lines=20, overlap_lines=3)
                all_chunks.extend(chunks)
                chunk_sources.extend([url] * len(chunks))

            except Exception as e:
                print(f"Failed to process {file_path}: {e}")

        self.document_chunks = all_chunks
        self.chunk_sources = chunk_sources

        print(f"Total chunks created: {len(all_chunks)}")

        # Generate embeddings
        if not self.embedding_model:
            self.load_embedding_model()

        if not self.document_chunks:
            logger.warning("No chunks to index")
            return

        print("Encoding chunks into embeddings...")

        # Use smaller batch size to prevent memory crashes
        batch_size = 16
        embeddings = []

        for i in tqdm(range(0, len(self.document_chunks), batch_size), desc="Encoding embeddings", unit="batch"):
            batch = self.document_chunks[i:i+batch_size]
            try:
                batch_embeddings = self.embedding_model.encode(batch, batch_size=batch_size, show_progress_bar=False)
                embeddings.extend(batch_embeddings)
                
                # Force garbage collection periodically
                if i % (batch_size * 50) == 0:
                    import gc
                    gc.collect()
                    
            except Exception as e:
                logger.warning(f"Failed to encode batch {i//batch_size}: {e}")

        # Save embeddings and chunks to cache
        if embeddings:
            print("Saving to cache for future runs...")
            try:
                with open(embeddings_cache_file, 'wb') as f:
                    pickle.dump(embeddings, f)
                with open(chunks_cache_file, 'wb') as f:
                    pickle.dump(self.document_chunks, f)
                with open(sources_cache_file, 'wb') as f:
                    pickle.dump(self.chunk_sources, f)
                print("Cache saved successfully")
            except Exception as e:
                print(f"Failed to save cache: {e}")

            # Build FAISS index
            self.build_faiss_index_from_embeddings(embeddings)
    def build_faiss_index_from_embeddings(self, embeddings):
        """Build FAISS index from pre-computed embeddings WITHOUT normalization"""
        try:
            print(f"Converting {len(embeddings)} embeddings to numpy array...")
            embeddings = np.array(embeddings).astype('float32')
            
            print(f"Embeddings shape: {embeddings.shape}")
            dimension = embeddings.shape[1]
            
            # Use L2 distance index instead of inner product - NO NORMALIZATION NEEDED
            print("Creating FAISS L2 index (no normalization required)...")
            self.index = faiss.IndexFlatL2(dimension)
            
            # Add embeddings directly without normalization
            print("Adding embeddings to index in batches...")
            batch_size = 1000
            
            for i in range(0, len(embeddings), batch_size):
                end_idx = min(i + batch_size, len(embeddings))
                batch = embeddings[i:end_idx]
                self.index.add(batch)
                
                if i % (batch_size * 10) == 0:
                    print(f"Added {end_idx:,}/{len(embeddings):,} embeddings to index...")
            
            print(f"✅ FAISS L2 index built with {len(embeddings)} chunks")
            
            # ========== ADD THIS BLOCK HERE ==========
            # Save sources to database
            if self.document_chunks and self.chunk_sources:
                print("Saving chunk sources to database...")
                import sqlite3
                conn = sqlite3.connect('terminology_rag.db')
                cursor = conn.cursor()
                
                # Add source_url column if doesn't exist
                try:
                    cursor.execute("ALTER TABLE documents ADD COLUMN source_url TEXT")
                    print("✓ Added source_url column")
                except sqlite3.OperationalError:
                    pass
                
                # Update rows with source URLs
                cursor.execute("SELECT id FROM documents ORDER BY id")
                doc_ids = [row[0] for row in cursor.fetchall()]
                
                for doc_id, source_url in zip(doc_ids, self.chunk_sources):
                    cursor.execute("UPDATE documents SET source_url = ? WHERE id = ?", (source_url, doc_id))
                
                conn.commit()
                conn.close()
                print(f"✓ Saved {len(self.chunk_sources)} source URLs to database")
            # ========== END BLOCK ==========
            
            # Save for reuse
            print("Saving index and metadata...")
            self.save_index_and_metadata()
            print("✅ Index saved successfully")
            
        except Exception as e:
            logger.error(f"❌ Failed to build FAISS index: {e}")
            import traceback
            traceback.print_exc()
    def extract_key_terms(self, query: str) -> List[str]:
        """Extract key terms from a query for better RAG search"""
        import re
        
        # Remove common words that don't help with terminology search
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
            'of', 'with', 'by', 'is', 'was', 'are', 'were', 'be', 'been', 'being',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 
            'should', 'may', 'might', 'can', 'must', 'shall', 'this', 'that',
            'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they',
            'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'her', 'its',
            'our', 'their'
        }
        
        # Clean and split the query
        words = re.findall(r'\b\w+\b', query.lower())
        
        # Remove stop words and short words
        key_terms = [word for word in words if word not in stop_words and len(word) > 2]
        
        # Also try to find multi-word phrases (2-3 words)
        phrases = []
        for i in range(len(key_terms) - 1):
            phrase = f"{key_terms[i]} {key_terms[i+1]}"
            phrases.append(phrase)
        
        return key_terms + phrases

    def search_relevant_content(self, query: str, k: int = 5) -> List[Dict]:
        """Enhanced search that returns chunks WITH source information"""
        if not self.embedding_model or not self.index:
            logger.error("RAG search failed: missing embedding model or index")
            return []

        try:
            query_lower = query.lower().strip()
            all_results = {}
            
            key_terms = self.extract_key_terms(query)
            
            # 1. Semantic search
            try:
                query_embedding = self.embedding_model.encode([query]).astype('float32')
                scores, indices = self.index.search(query_embedding, k * 2)
                
                for score, idx in zip(scores[0], indices[0]):
                    if idx < len(self.document_chunks):
                        chunk = self.document_chunks[idx]
                        source_url = self.chunk_sources[idx] if idx < len(self.chunk_sources) else "Unknown"
                        similarity = 1.0 / (1.0 + float(score))
                        
                        chunk_key = f"{idx}_{chunk[:50]}"
                        if chunk_key not in all_results or all_results[chunk_key]['score'] < similarity:
                            all_results[chunk_key] = {
                                'text': chunk,
                                'score': similarity,
                                'source': source_url,
                                'index': idx
                            }

            except Exception as e:
                logger.error(f"Semantic search failed: {e}")

            # 2. Search for key terms
            for term in key_terms:
                if len(term) > 3:
                    try:
                        term_embedding = self.embedding_model.encode([term]).astype('float32')
                        scores, indices = self.index.search(term_embedding, k)
                        
                        for score, idx in zip(scores[0], indices[0]):
                            if idx < len(self.document_chunks):
                                chunk = self.document_chunks[idx]
                                source_url = self.chunk_sources[idx] if idx < len(self.chunk_sources) else "Unknown"
                                similarity = 1.0 / (1.0 + float(score))
                                
                                chunk_key = f"{idx}_{chunk[:50]}"
                                if chunk_key not in all_results or all_results[chunk_key]['score'] < similarity * 1.1:
                                    all_results[chunk_key] = {
                                        'text': chunk,
                                        'score': similarity * 1.1,
                                        'source': source_url,
                                        'index': idx
                                    }
                                
                    except Exception as e:
                        logger.warning(f"Term search failed for '{term}': {e}")

            # 3. Exact matching
            for i, chunk in enumerate(self.document_chunks):
                chunk_lower = chunk.lower()
                source_url = self.chunk_sources[i] if i < len(self.chunk_sources) else "Unknown"
                
                if query_lower in chunk_lower:
                    chunk_key = f"{i}_{chunk[:50]}"
                    if chunk_key not in all_results or all_results[chunk_key]['score'] < 0.95:
                        all_results[chunk_key] = {
                            'text': chunk,
                            'score': 0.95,
                            'source': source_url,
                            'index': i
                        }
                
                for term in key_terms:
                    if term in chunk_lower:
                        chunk_key = f"{i}_{chunk[:50]}"
                        term_score = 0.8 + (len(term.split()) * 0.05)
                        if chunk_key not in all_results or all_results[chunk_key]['score'] < term_score:
                            all_results[chunk_key] = {
                                'text': chunk,
                                'score': term_score,
                                'source': source_url,
                                'index': i
                            }

            final_results = sorted(all_results.values(), key=lambda item: item['score'], reverse=True)
            logger.info(f"Search for '{query}' found {len(final_results)} results with sources")
            return final_results[:k]

        except Exception as e:
            logger.error(f"Search failed: {e}")
            return []
    
    def add_documents(self, documents: List[str], sources: List[Dict]):
        """Add new documents to the RAG index dynamically"""
        try:
            if not self.embedding_model:
                logger.warning("Embedding model not loaded, loading now...")
                self.load_embedding_model()
            
            if not self.embedding_model:
                logger.error("Cannot add documents: embedding model failed to load")
                return False
            
            logger.info(f"Adding {len(documents)} documents to RAG index...")
            
            # Process and chunk documents
            new_chunks = []
            new_sources = []
            
            for doc, source in zip(documents, sources):
                # Simple chunking for glossary terms (they're already small)
                if len(doc) < 500:
                    new_chunks.append(doc)
                    source_url = source.get('url', 'custom_upload')
                    new_sources.append(source_url)
                else:
                    # Chunk larger documents
                    chunks = self.chunk_text_by_lines(doc, source.get('url', 'custom_upload'))
                    new_chunks.extend(chunks)
                    new_sources.extend([source.get('url', 'custom_upload')] * len(chunks))
            
            # Add to existing collections
            self.document_chunks.extend(new_chunks)
            self.chunk_sources.extend(new_sources)
            
            # Rebuild FAISS index with all documents
            logger.info(f"Rebuilding FAISS index with {len(self.document_chunks)} total chunks...")
            embeddings = self.embedding_model.encode(self.document_chunks, show_progress_bar=True)
            embeddings = embeddings.astype('float32')
            
            # Create new index
            dimension = embeddings.shape[1]
            self.index = faiss.IndexFlatL2(dimension)
            self.index.add(embeddings)
            
            logger.info(f"✅ Successfully added {len(new_chunks)} new chunks to RAG index")
            logger.info(f"✅ Total chunks in index: {len(self.document_chunks)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Failed to add documents to RAG: {e}")
            return False

class TerminologyAwareTranslator:
    """Enhanced translator with RAG-powered terminology awareness"""
    
    def __init__(self):
        global rag_system
        self.rag_system = rag_system
    
    def filter_chunk_for_relevance(self, chunk_text: str, query: str, query_terms: set) -> str:
        """
        Filter chunk to only show lines relevant to the query.
        Prevents showing adjacent unrelated terms from terminology dictionaries.
        """
        try:
            lines = chunk_text.split('\n')
            relevant_lines = []
            
            # Normalize query for better matching
            query_lower = query.lower()
            query_normalized = unicodedata.normalize('NFKC', query_lower)
            
            # Extract key terms from query (words longer than 3 chars)
            key_terms = [term for term in query_terms if len(term) > 3]
            
            # Strategy: Keep lines that contain query terms or are contextually related
            for i, line in enumerate(lines):
                line_lower = line.lower()
                line_normalized = unicodedata.normalize('NFKC', line_lower)
                
                # Keep source headers and document context
                if line.startswith('SOURCE:') or line.startswith('DOCUMENT START:') or line.startswith('---'):
                    relevant_lines.append(line)
                    continue
                
                # Skip empty lines (but might add back for formatting)
                if not line.strip():
                    continue
                
                # Check if line contains any query term
                contains_query_term = False
                for term in key_terms:
                    term_normalized = unicodedata.normalize('NFKC', term.lower())
                    if term_normalized in line_normalized:
                        contains_query_term = True
                        break
                
                # Also check for exact query substring
                if query_normalized in line_normalized:
                    contains_query_term = True
                
                if contains_query_term:
                    # DICTIONARY FIX: Keep matching line + next line (which has translation)
                    if line not in relevant_lines:
                        relevant_lines.append(line)
                    # Add next line ONLY if it contains Greek characters (is a translation, not another English term)
                    if i + 1 < len(lines) and lines[i + 1].strip():
                        next_line = lines[i + 1]
                        # Check if next line has Greek characters
                        has_greek = bool(re.search(r'[Α-Ωα-ωά-ώ]', next_line))
                        if has_greek and next_line not in relevant_lines:
                            relevant_lines.append(next_line)
            
            # If we filtered out everything, return original chunk
            if len(relevant_lines) == 0:
                logger.warning(f"Chunk filter removed everything, returning original")
                return chunk_text
            
            # For dictionary entries, even 1-2 lines is fine
            # Only return original if we got nothing useful
            
            filtered = '\n'.join(relevant_lines)
            
            # Log filtering statistics
            original_lines = len(lines)
            filtered_lines = len(relevant_lines)
            if filtered_lines < original_lines * 0.8:
                logger.info(f"📝 Filtered chunk: {original_lines} → {filtered_lines} lines")
            
            return filtered
            
        except Exception as e:
            logger.warning(f"Chunk filtering failed: {e}, returning original")
            return chunk_text
        
    async def translate_with_rag(self, text: str, source_lang: str, target_lang: str, 
                            style: str = "formal", provider: str = None, db = None, session_id: Optional[str] = None) -> Dict:
        """Enhanced translation with glossary and source-aware chunks"""
        
        # Check custom glossary FIRST (check BOTH directions)
        glossary_terms = []
        if db:
            try:
                logger.info(f"🔍 Searching glossary for: '{text}' ({source_lang} → {target_lang})")
                
                # Forward direction: source_lang -> target_lang
                glossary_forward = db.query(CustomGlossary).filter(
                    CustomGlossary.source_language == source_lang,
                    CustomGlossary.target_language == target_lang
                ).all()
                
                # Reverse direction: target_lang -> source_lang (use in reverse!)
                glossary_reverse = db.query(CustomGlossary).filter(
                    CustomGlossary.source_language == target_lang,
                    CustomGlossary.target_language == source_lang
                ).all()
                
                logger.info(f"📖 Found {len(glossary_forward)} forward entries, {len(glossary_reverse)} reverse entries")
                
                text_lower = text.lower()
                # Normalize text for better Greek matching
                import unicodedata
                text_normalized = unicodedata.normalize('NFKC', text_lower)
                
                # Check forward matches (normal direction)
                for term in glossary_forward:
                    source_normalized = unicodedata.normalize('NFKC', term.source_term.lower())
                    if source_normalized in text_normalized:
                        # Include context/domain information for LLM to make informed decisions
                        if term.context:
                            glossary_terms.append(f"{term.source_term} → {term.target_term} [{term.context}]")
                        else:
                            glossary_terms.append(f"{term.source_term} → {term.target_term}")
                        logger.info(f"✅ Glossary match (forward): {term.source_term} → {term.target_term}")
                
                # Check reverse matches (use backwards!)
                for term in glossary_reverse:
                    # For reverse: if translating el->en but glossary has en->el,
                    # check if Greek target_term appears in text, then use source_term as translation
                    target_normalized = unicodedata.normalize('NFKC', term.target_term.lower())
                    if target_normalized in text_normalized:
                        # Reversed: target → source (e.g., έμβαση → onset)
                        # Include context/domain information for LLM to make informed decisions
                        if term.context:
                            glossary_terms.append(f"{term.target_term} → {term.source_term} [{term.context}]")
                        else:
                            glossary_terms.append(f"{term.target_term} → {term.source_term}")
                        logger.info(f"✅ Glossary match (reverse): {term.target_term} → {term.source_term}")
                
                if len(glossary_terms) == 0:
                    logger.warning(f"⚠️ No glossary matches found. Text: '{text[:100]}...'")
                else:
                    logger.info(f"✨ Found {len(glossary_terms)} glossary matches total!")
                
            except Exception as e:
                logger.warning(f"Glossary lookup failed: {e}")
        
        relevant_chunks = self.rag_system.search_relevant_content(text, k=8)
        
        context_parts = []
        sources_used = []
        
        # Add glossary terms FIRST
        if glossary_terms:
            context_parts.append("=== CUSTOM GLOSSARY (USE THESE TRANSLATIONS) ===")
            context_parts.extend(glossary_terms)
            context_parts.append("")
        
        # Extract query terms for relevance filtering
        query_terms = set(text.lower().split())
        
        for i, chunk_data in enumerate(relevant_chunks[:5], 1):
            chunk_text = chunk_data['text']
            source_url = chunk_data['source']
            score = chunk_data['score']
            
            # Filter chunk to only show relevant lines
            filtered_chunk = self.filter_chunk_for_relevance(chunk_text, text, query_terms)
            
            if not filtered_chunk:
                continue  # Skip if nothing relevant
            
            doc_name = source_url.split('/')[-1] if '/' in source_url else source_url
            
            # Add source info WITH document name clearly labeled
            context_parts.append(f"=== SOURCE {i}: {doc_name} ===")
            context_parts.append(f"Document: {doc_name}")
            context_parts.append(f"URL: {source_url}")
            context_parts.append(f"Relevance: {score:.2f}")
            context_parts.append("---")
            context_parts.append(filtered_chunk.strip())
            context_parts.append("")
            
            sources_used.append({
                'source': doc_name,
                'url': source_url,
                'score': score,
                'index': i
            })
        
        context = "\n".join(context_parts)
        
        if source_lang == 'el' and target_lang == 'en':
            prompt = f"""You are a professional Greek-to-English translator specializing in context-aware translation.

    RELEVANT TERMINOLOGY SOURCES:
    {context}

    TEXT TO TRANSLATE: {text}

    CRITICAL INSTRUCTIONS:
    1. **CONTEXT-AWARE TRANSLATION**:
       - First, identify the domain/subject matter of the text (linguistics, technology, medicine, sports, etc.)
       - Choose terminology that fits the SPECIFIC CONTEXT and usage in this text
       - Consider the full sentence meaning, not just word-by-word translation
       - If a term has multiple valid translations, choose the one that fits THIS context best
    
    2. **CUSTOM GLOSSARY - DOMAIN-AWARE PRIORITY**:
       - Custom glossary terms are shown with [domain/context] information
       - **USE glossary term IF**: Its domain matches your identified text domain, OR it has no domain specified
       - **DO NOT use glossary term IF**: Its domain clearly doesn't match the text domain
       - Example: Don't use "run → εκτελεί [Domain: computing]" if translating a sports text
       - If no domain specified in glossary → assume it's general-purpose and use it
    
    3. **RAG SOURCES - DOMAIN-AWARE SELECTION**:
       - ONLY use terms from sources that EXACTLY MATCH the words in the text being translated
       - Consider the domain/context of the source - prefer sources from the same domain as your text
       - **DO NOT use terms from sources that don't appear in the text** - even if nearby in the source
    
    4. **SEMANTIC CHECK**: Verify the translation makes sense in the FULL SENTENCE context. If a term seems unrelated, DON'T use it.
    
    5. Example: If translating "ascender", DO NOT use "blinking" translations even if they appear in the same source document
    
    6. **Maintain {style} style** while adapting to the text's domain and register
    
    7. When citing sources:
       - Format: [SOURCE X: document_name.pdf]
       - Or: [From document_name.pdf]
       - ALWAYS include the document name from the SOURCE header
       - Example: "αναβάτης [SOURCE 1: terminology_dict.pdf]"
    
    8. If a source shows multiple unrelated terms, use ONLY the one that matches your text

    Translation:"""

        else:
            prompt = f"""You are a professional English-to-Greek translator specializing in context-aware translation.

    RELEVANT TERMINOLOGY SOURCES:
    {context}

    TEXT TO TRANSLATE: {text}

    CRITICAL INSTRUCTIONS:
    1. **CONTEXT-AWARE TRANSLATION**:
       - First, identify the domain/subject matter of the text (linguistics, technology, medicine, sports, etc.)
       - Choose terminology that fits the SPECIFIC CONTEXT and usage in this text
       - Consider the full sentence meaning, not just word-by-word translation
       - If a term has multiple valid translations, choose the one that fits THIS context best
    
    2. **CUSTOM GLOSSARY - DOMAIN-AWARE PRIORITY**:
       - Custom glossary terms are shown with [domain/context] information
       - **USE glossary term IF**: Its domain matches your identified text domain, OR it has no domain specified
       - **DO NOT use glossary term IF**: Its domain clearly doesn't match the text domain
       - Example: Don't use "run → εκτελεί [Domain: computing]" if translating a sports text
       - If no domain specified in glossary → assume it's general-purpose and use it
    
    3. **RAG SOURCES - DOMAIN-AWARE SELECTION**:
       - ONLY use terms from sources that EXACTLY MATCH the words in the text being translated
       - Consider the domain/context of the source - prefer sources from the same domain as your text
       - **DO NOT use terms from sources that don't appear in the text** - even if nearby in the source
    
    4. **SEMANTIC CHECK**: Verify the translation makes sense in the FULL SENTENCE context. If a term seems unrelated, DON'T use it.
    
    5. **Maintain {style} style** while adapting to the text's domain and register
    
    6. When citing sources:
       - Format: [SOURCE X: document_name.pdf]
       - Or: [From document_name.pdf]
       - ALWAYS include the document name from the SOURCE header
       - Example: "αναβάτης [SOURCE 1: terminology_dict.pdf]"
    
    7. If a source shows multiple unrelated terms, use ONLY the one that matches your text

    Translation:"""

        try:
            translation = await generate_with_timeout_multi(prompt, provider, timeout=60, session_id=session_id)
            
            # Parse citations from the translation
            import re
            citations = re.findall(r'\[SOURCE (\d+)\]', translation)
            cited_indices = list(set([int(i) for i in citations]))
            
            # Filter to only sources that were actually cited
            actually_used_sources = [s for s in sources_used if s['index'] in cited_indices]
            
            # If no citations found, assume all sources might have been used implicitly
            if not actually_used_sources and sources_used:
                actually_used_sources = sources_used[:2]  # Show top 2 as potentially used
            
            return {
                'translated_text': translation.strip(),
                'source_language': source_lang,
                'target_language': target_lang,
                'style': style,
                'sources_used': len(actually_used_sources),
                'sources': actually_used_sources,
                'llm_provider': provider or llm_manager.default_provider.value,
                'confidence_score': sum(s['score'] for s in actually_used_sources) / len(actually_used_sources) if actually_used_sources else 0.5
            }
            
        except Exception as e:
            logger.error(f"Translation failed: {e}")
            raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

    async def chat_about_terminology(self, user_message: str, conversation_history: List[Dict] = None, provider: str = None, db = None, session_id: Optional[str] = None) -> Dict:
        """Chat about terminology with glossary and RAG enhancement"""
        
        # Check custom glossary FIRST
        glossary_terms = []
        if db:
            try:
                # Get all glossary terms (we don't know language direction in chat)
                glossary_matches = db.query(CustomGlossary).all()
                
                for term in glossary_matches:
                    # Check if either source or target term appears in the message
                    if (term.source_term.lower() in user_message.lower() or 
                        term.target_term.lower() in user_message.lower()):
                        # Include context/domain information for LLM to make informed decisions
                        if term.context:
                            glossary_terms.append(f"{term.source_term} ({term.source_language}) ↔ {term.target_term} ({term.target_language}) [{term.context}]")
                        else:
                            glossary_terms.append(f"{term.source_term} ({term.source_language}) ↔ {term.target_term} ({term.target_language})")
            except Exception as e:
                logger.warning(f"Glossary lookup failed: {e}")
        
        # Extract the actual term if the user is asking about a specific term
        # Common patterns: "tell me about X", "what is X", "πες μου για X", "τι είναι X", etc.
        search_query = user_message
        
        # Check for quoted terms
        import re
        quoted_terms = re.findall(r'["\']([^"\']+)["\']', user_message)
        if quoted_terms:
            # If user put term in quotes, search for that specifically
            search_query = quoted_terms[0]
            logger.info(f"📝 Extracted quoted term for search: '{search_query}'")
        else:
            # Check for common question patterns
            term_patterns = [
                r'(?:tell me about|what is|define|explain|translate)\s+(?:the term\s+)?["\']?([^"\'?\n]+?)["\']?(?:\?|$)',
                r'(?:πες μου για|τι είναι|ορισμός|εξήγησε|μετάφρασε)\s+(?:τον όρο\s+)?["\']?([^"\'?\n]+?)["\']?(?:\?|$)',
            ]
            for pattern in term_patterns:
                match = re.search(pattern, user_message, re.IGNORECASE)
                if match:
                    search_query = match.group(1).strip()
                    logger.info(f"📝 Extracted term from question: '{search_query}'")
                    break
        
        relevant_chunks = self.rag_system.search_relevant_content(search_query, k=8)
        
        # Build context with source tracking
        context_parts = []
        sources_list = []
        query_terms = set(search_query.lower().split())
        
        # Add glossary terms FIRST
        if glossary_terms:
            context_parts.append("=== CUSTOM GLOSSARY (USE THESE TRANSLATIONS) ===")
            context_parts.extend(glossary_terms)
            context_parts.append("")
        
        for i, chunk_data in enumerate(relevant_chunks[:5], 1):
            chunk_text = chunk_data['text']
            source_url = chunk_data['source']
            score = chunk_data['score']
            
            # Filter chunk to only show relevant lines
            filtered_chunk = self.filter_chunk_for_relevance(chunk_text, user_message, query_terms)
            
            if not filtered_chunk:
                continue  # Skip if nothing relevant
            
            doc_name = source_url.split('/')[-1] if '/' in source_url else source_url
            
            # Add source info WITH document name clearly labeled
            context_parts.append(f"=== SOURCE {i}: {doc_name} ===")
            context_parts.append(f"Document: {doc_name}")
            context_parts.append(f"URL: {source_url}")
            context_parts.append(f"Relevance: {score:.2f}")
            context_parts.append("---")
            context_parts.append(filtered_chunk.strip())
            context_parts.append("")
            
            sources_list.append({
                'source': doc_name,
                'url': source_url,
                'score': score,
                'index': i
            })
        
        context = "\n".join(context_parts)
        
        history_text = ""
        if conversation_history:
            for msg in conversation_history[-5:]:
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                history_text += f"{role.title()}: {content}\n"
                
        prompt = f"""You are a terminology expert with access to authoritative Greek-English translation resources.

    RELEVANT TERMINOLOGY SOURCES:
    {context}

    CONVERSATION HISTORY:
    {history_text}

    USER QUESTION: {user_message}

    CRITICAL INSTRUCTIONS:
    1. **CONTEXT-AWARE RESPONSE**:
       - Identify the domain/subject matter of the question (linguistics, technology, medicine, etc.)
       - Provide terminology that fits the SPECIFIC CONTEXT
       - If a term has multiple meanings/translations, specify which domain each applies to
    
    2. **CUSTOM GLOSSARY - DOMAIN-AWARE PRIORITY**:
       - Custom glossary terms are shown with [domain/context] information
       - **USE glossary term IF**: Its domain matches the question's domain, OR it has no domain specified
       - **DO NOT use glossary term IF**: Its domain clearly doesn't match
       - If multiple glossary entries exist for the same term in different domains, mention all relevant ones
    
    3. **YOU MUST PROVIDE A CLEAR ANSWER** - don't just list sources
    
    4. From sources: ONLY use terms that EXACTLY MATCH the words in the question
    
    5. **DO NOT use terms from sources that don't appear in the question** - even if they appear nearby in the same source
    
    6. **SEMANTIC CHECK**: Verify the translation makes sense in context. If a term seems unrelated, DON'T mention it.
    
    7. Example: If asked about "ascender", DO NOT mention "blinking" or "αναβόσβημα" even if they appear in the same source
    
    8. When citing sources:
       - Format: [SOURCE X: document_name.pdf]
       - Or: [From document_name.pdf]
       - ALWAYS include the document name from the SOURCE header
       - Example: "αναβάτης [SOURCE 1: terminology_dict.pdf]"
    
    9. If multiple valid translations exist for THE SAME TERM, list them. But DON'T list translations for different terms.
    
    10. Format answer as: "The translation for X is Y [SOURCE N: document.pdf]"
    
    11. Be direct and helpful

    Response:"""

        try:
            response = await generate_with_timeout_multi(prompt, provider, timeout=60, session_id=session_id)
            
            # Parse citations from the response
            import re
            citations = re.findall(r'\[SOURCE (\d+)\]', response)
            cited_indices = list(set([int(i) for i in citations]))
            
            # Filter to only sources that were actually cited
            actually_used_sources = [s for s in sources_list if s['index'] in cited_indices]
            
            # If no citations found, assume all sources might have been used implicitly
            if not actually_used_sources and sources_list:
                actually_used_sources = sources_list[:2]  # Show top 2 as potentially used
            
            return {
                'response': response.strip(),
                'sources': actually_used_sources
            }
            
        except Exception as e:
            logger.error(f"Chat failed: {e}")
            return {
                'response': "I'm sorry, I encountered an error while processing your question. Please try again.",
                'sources': []
            }
class TerminologyChat:
    """Chat interface for terminology discussions"""
    
    def __init__(self):
        global rag_system
        self.rag_system = rag_system
        # Use the translator's filter method
        global terminology_translator
        self.translator = terminology_translator
    
    def filter_chunk_for_relevance(self, chunk_text: str, query: str, query_terms: set) -> str:
        # USE THE WORKING TRANSLATOR FILTER
        if self.translator:
            return self.translator.filter_chunk_for_relevance(chunk_text, query, query_terms)
        # Fallback to local implementation
        return self._filter_chunk_fallback(chunk_text, query, query_terms)
    
    def _filter_chunk_fallback(self, chunk_text: str, query: str, query_terms: set) -> str:
        """
        Filter chunk to only show lines relevant to the query.
        Prevents showing adjacent unrelated terms from terminology dictionaries.
        """
        try:
            lines = chunk_text.split('\n')
            relevant_lines = []
            
            # Normalize query for better matching
            query_lower = query.lower()
            query_normalized = unicodedata.normalize('NFKC', query_lower)
            
            # Extract key terms from query (words longer than 3 chars)
            key_terms = [term for term in query_terms if len(term) > 3]
            
            # Strategy: Keep lines that contain query terms or are contextually related
            for i, line in enumerate(lines):
                line_lower = line.lower()
                line_normalized = unicodedata.normalize('NFKC', line_lower)
                
                # Keep source headers and document context
                if line.startswith('SOURCE:') or line.startswith('DOCUMENT START:') or line.startswith('---'):
                    relevant_lines.append(line)
                    continue
                
                # Skip empty lines (but might add back for formatting)
                if not line.strip():
                    continue
                
                # Check if line contains any query term
                contains_query_term = False
                for term in key_terms:
                    term_normalized = unicodedata.normalize('NFKC', term.lower())
                    if term_normalized in line_normalized:
                        contains_query_term = True
                        break
                
                # Also check for exact query substring
                if query_normalized in line_normalized:
                    contains_query_term = True
                
                if contains_query_term:
                    # DICTIONARY FIX: Keep matching line + next line (which has translation)
                    if line not in relevant_lines:
                        relevant_lines.append(line)
                    # Add next line ONLY if it contains Greek characters (is a translation, not another English term)
                    if i + 1 < len(lines) and lines[i + 1].strip():
                        next_line = lines[i + 1]
                        # Check if next line has Greek characters
                        has_greek = bool(re.search(r'[Α-Ωα-ωά-ώ]', next_line))
                        if has_greek and next_line not in relevant_lines:
                            relevant_lines.append(next_line)
            
            # If we filtered out everything, return original chunk
            if len(relevant_lines) == 0:
                logger.warning(f"Chunk filter removed everything, returning original")
                return chunk_text
            
            # For dictionary entries, even 1-2 lines is fine
            # Only return original if we got nothing useful
            
            filtered = '\n'.join(relevant_lines)
            
            # Log filtering statistics
            original_lines = len(lines)
            filtered_lines = len(relevant_lines)
            if filtered_lines < original_lines * 0.8:
                logger.info(f"📝 Chat chunk filtered: {original_lines} → {filtered_lines} lines")
            
            return filtered
            
        except Exception as e:
            logger.warning(f"Chunk filtering failed: {e}, returning original")
            return chunk_text
         
    async def chat_about_terminology(self, user_message: str, conversation_history: List[Dict] = None, provider: str = None, db = None, session_id: Optional[str] = None) -> Dict:
        """Chat about terminology with glossary and RAG enhancement"""
        
        # Check custom glossary FIRST
        glossary_terms = []
        if db:
            try:
                # Get all glossary terms (we don't know language direction in chat)
                glossary_matches = db.query(CustomGlossary).all()
                
                for term in glossary_matches:
                    # Check if either source or target term appears in the message
                    if (term.source_term.lower() in user_message.lower() or 
                        term.target_term.lower() in user_message.lower()):
                        # Include context/domain information for LLM to make informed decisions
                        if term.context:
                            glossary_terms.append(f"{term.source_term} ({term.source_language}) ↔ {term.target_term} ({term.target_language}) [{term.context}]")
                        else:
                            glossary_terms.append(f"{term.source_term} ({term.source_language}) ↔ {term.target_term} ({term.target_language})")
            except Exception as e:
                logger.warning(f"Glossary lookup failed: {e}")
        
        # Extract the actual term if the user is asking about a specific term
        # Common patterns: "tell me about X", "what is X", "πες μου για X", "τι είναι X", etc.
        search_query = user_message
        
        # Check for quoted terms
        import re
        quoted_terms = re.findall(r'["\']([^"\']+)["\']', user_message)
        if quoted_terms:
            # If user put term in quotes, search for that specifically
            search_query = quoted_terms[0]
            logger.info(f"📝 Extracted quoted term for search: '{search_query}'")
        else:
            # Check for common question patterns
            term_patterns = [
                r'(?:tell me about|what is|define|explain|translate)\s+(?:the term\s+)?["\']?([^"\'?\n]+?)["\']?(?:\?|$)',
                r'(?:πες μου για|τι είναι|ορισμός|εξήγησε|μετάφρασε)\s+(?:τον όρο\s+)?["\']?([^"\'?\n]+?)["\']?(?:\?|$)',
            ]
            for pattern in term_patterns:
                match = re.search(pattern, user_message, re.IGNORECASE)
                if match:
                    search_query = match.group(1).strip()
                    logger.info(f"📝 Extracted term from question: '{search_query}'")
                    break
        
        relevant_chunks = self.rag_system.search_relevant_content(search_query, k=8)
        
        # Build context with source tracking
        context_parts = []
        sources_list = []
        query_terms = set(search_query.lower().split())
        
        # Add glossary terms FIRST
        if glossary_terms:
            context_parts.append("=== CUSTOM GLOSSARY (USE THESE TRANSLATIONS) ===")
            context_parts.extend(glossary_terms)
            context_parts.append("")
        
        for i, chunk_data in enumerate(relevant_chunks[:5], 1):
            chunk_text = chunk_data['text']
            source_url = chunk_data['source']
            score = chunk_data['score']
            
            # Filter chunk to only show relevant lines
            filtered_chunk = self.filter_chunk_for_relevance(chunk_text, search_query, query_terms)
            
            if not filtered_chunk:
                continue  # Skip if nothing relevant
            
            doc_name = source_url.split('/')[-1] if '/' in source_url else source_url
            
            # Add source info WITH document name clearly labeled
            context_parts.append(f"=== SOURCE {i}: {doc_name} ===")
            context_parts.append(f"Document: {doc_name}")
            context_parts.append(f"URL: {source_url}")
            context_parts.append(f"Relevance: {score:.2f}")
            context_parts.append("---")
            context_parts.append(filtered_chunk.strip())
            context_parts.append("")
            
            sources_list.append({
                'source': doc_name,
                'url': source_url,
                'score': score,
                'index': i
            })
        
        context = "\n".join(context_parts)
        
        history_text = ""
        if conversation_history:
            for msg in conversation_history[-5:]:
                role = msg.get('role', 'user')
                content = msg.get('content', '')
                history_text += f"{role.title()}: {content}\n"
                
        prompt = f"""You are a terminology expert with access to authoritative Greek-English translation resources.

    RELEVANT TERMINOLOGY SOURCES:
    {context}

    CONVERSATION HISTORY:
    {history_text}

    USER QUESTION: {user_message}

    CRITICAL INSTRUCTIONS:
    1. **CONTEXT-AWARE RESPONSE**:
       - Identify the domain/subject matter of the question (linguistics, technology, medicine, etc.)
       - Provide terminology that fits the SPECIFIC CONTEXT
       - If a term has multiple meanings/translations, specify which domain each applies to
    
    2. **CUSTOM GLOSSARY - DOMAIN-AWARE PRIORITY**:
       - Custom glossary terms are shown with [domain/context] information
       - **USE glossary term IF**: Its domain matches the question's domain, OR it has no domain specified
       - **DO NOT use glossary term IF**: Its domain clearly doesn't match
       - If multiple glossary entries exist for the same term in different domains, mention all relevant ones
    
    3. **YOU MUST PROVIDE A CLEAR ANSWER** - don't just list sources
    
    4. From sources: ONLY use terms that EXACTLY MATCH the words in the question
    
    5. **DO NOT use terms from sources that don't appear in the question** - even if they appear nearby in the same source
    
    6. **SEMANTIC CHECK**: Verify the translation makes sense in context. If a term seems unrelated, DON'T mention it.
    
    7. Example: If asked about "ascender", DO NOT mention "blinking" or "αναβόσβημα" even if they appear in the same source
    
    8. When citing sources:
       - Format: [SOURCE X: document_name.pdf]
       - Or: [From document_name.pdf]
       - ALWAYS include the document name from the SOURCE header
       - Example: "αναβάτης [SOURCE 1: terminology_dict.pdf]"
    
    9. If multiple valid translations exist for THE SAME TERM, list them. But DON'T list translations for different terms.
    
    10. Format answer as: "The translation for X is Y [SOURCE N: document.pdf]"
    
    11. Be direct and helpful

    Response:"""

        try:
            response = await generate_with_timeout_multi(prompt, provider, timeout=60, session_id=session_id)
            
            # Parse citations from the response
            import re
            citations = re.findall(r'\[SOURCE (\d+)\]', response)
            cited_indices = list(set([int(i) for i in citations]))
            
            # Filter to only sources that were actually cited
            actually_used_sources = [s for s in sources_list if s['index'] in cited_indices]
            
            # If no citations found, assume all sources might have been used implicitly
            if not actually_used_sources and sources_list:
                actually_used_sources = sources_list[:2]  # Show top 2 as potentially used
            
            return {
                'response': response.strip(),
                'sources': actually_used_sources
            }
            
        except Exception as e:
            logger.error(f"Chat failed: {e}")
            return {
                'response': "I'm sorry, I encountered an error while processing your question. Please try again.",
                'sources': []
            }
    def build_chat_context(self, relevant_chunks: List[Dict]) -> str:
        """Build context for chat"""
        if not relevant_chunks:
            return "No specific terminology references found for this topic."
            
        context_parts = []
        for chunk_data in relevant_chunks:
            score = chunk_data.get('score', 0.0)
            if score > 0.5:
                chunk_text = chunk_data.get('text', '')
                cleaned_chunk = chunk_text.strip()
                if len(cleaned_chunk) > 30:
                    context_parts.append(cleaned_chunk)
                    
        return "\n---\n".join(context_parts[:3])

# Global instances
scraper = EletoDocumentScraper()
rag_system = TerminologyRAGSystem()
terminology_translator = None
terminology_chat = None

def initialize_terminology_system():
    global terminology_translator, terminology_chat
    terminology_translator = TerminologyAwareTranslator()
    terminology_chat = TerminologyChat()

    rag_system.load_embedding_model()

    if not rag_system.load_existing_index():
        logger.info("⚠️ No saved index found. Rebuilding from files...")
        rag_system.load_and_process_saved_files()
    else:
        logger.info("✅ Using saved FAISS index and metadata")

async def scrape_eleto_documents():
    try:
        scraped_content = await scraper.scrape_everything()
        
        if scraped_content:
            # Step 1: Save all files
            saved_count = rag_system.save_files_only(scraped_content)
            print(f"Saved {saved_count} files to disk")
            
            # Step 2: Process them (you can do this later if needed)
            print("Starting chunk processing and embedding generation...")
            rag_system.load_and_process_saved_files()
            
            return saved_count
    except Exception as e:
        logger.error(f"Scraping failed: {e}")
        return 0
    finally:
        await scraper.close_session()
# Language detection function
def detect_text_language(text: str) -> str:
    """Detect the language of the input text"""
    try:
        greek_chars = len(re.findall(r'[\u0370-\u03FF\u1F00-\u1FFF]', text))
        total_letters = len(re.findall(r'[a-zA-ZάέήίόύώαβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]', text))
        
        if total_letters == 0:
            return "en"
        
        greek_ratio = greek_chars / total_letters
        
        if greek_ratio > 0.3:
            logger.info(f"Detected Greek text: {greek_ratio:.2%} Greek characters")
            return "el"
        
        greek_indicators = [
            'και', 'που', 'είναι', 'από', 'για', 'στο', 'στη', 'στην', 'του', 'της', 'των',
            'με', 'αυτό', 'αυτή', 'αυτός', 'όλα', 'όλες', 'όλοι', 'μου', 'σου', 'του',
            'μας', 'σας', 'τους', 'θα', 'να', 'ο', 'η', 'το', 'οι', 'τα'
        ]
        
        words = text.lower().split()
        greek_word_count = sum(1 for word in words if any(indicator in word for indicator in greek_indicators))
        
        if len(words) > 0 and (greek_word_count / len(words)) > 0.1:
            logger.info(f"Detected Greek text: {greek_word_count}/{len(words)} Greek words")
            return "el"
        
        return "en"
        
    except Exception as e:
        logger.warning(f"Language detection failed: {e}")
        return "en"

# Placeholder functions for compatibility
def check_per_request_word_limit(text: str, user_tier: str):
    word_count = len(text.split())
    per_request_limits = {
        "free": 800,
        "premium": 5000,
        "enterprise": 15000
    }
    limit = per_request_limits.get(user_tier, 800)
    if word_count > limit:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Text too long: {word_count} words. Your {user_tier} plan allows up to {limit} words per request."
        )

def check_daily_word_limit(db: SQLASessionType, user: User, text: str):
    # Placeholder - implement your word limit logic
    pass

def record_word_usage(db: SQLASessionType, user: User, text: str):
    # Placeholder - implement your word usage tracking
    pass

def get_current_user(request: Request, db: SQLASessionType) -> User:
    # Placeholder - implement your authentication
    # Return a mock user for now
    pass

# FASTAPI ENDPOINTS

def get_session_id(request: Request) -> str:
    """Get or create session ID from request"""
    session_id = request.cookies.get("session_id")
    if not session_id:
        # Create a new session ID based on client IP and user agent
        import hashlib
        client_info = f"{request.client.host if request.client else 'unknown'}-{request.headers.get('user-agent', 'unknown')}"
        session_id = hashlib.md5(client_info.encode()).hexdigest()
    return session_id

@app.get("/api/llm-models")
async def get_llm_models(request: Request):
    """Get all available LLM models with their availability status"""
    session_id = get_session_id(request)
    models = llm_manager.get_available_models(session_id)
    
    # Group by provider for easier frontend handling
    grouped = {
        "gemini": [],
        "anthropic": [],
        "openai": []
    }
    
    for model_id, info in models.items():
        provider = info["provider"]
        grouped[provider].append({
            "id": model_id,
            "name": info["name"],
            "description": info["description"],
            "available": info["available"],
            "has_user_key": info["has_user_key"],
            "max_tokens": info["max_tokens"],
            "supports_streaming": info["supports_streaming"]
        })
    
    return JSONResponse(content={
        "models": models,
        "grouped": grouped,
        "default": llm_manager.default_provider.value
    })

@app.get("/api/llm-providers")
async def get_llm_providers(request: Request):
    """Legacy endpoint - redirects to new models endpoint"""
    return await get_llm_models(request)

@app.post("/api/set-api-key")
async def set_api_key(
    request: Request,
    provider_type: str = Form(...),  # "gemini", "anthropic", or "openai"
    api_key: str = Form(...)
):
    """Set user API key for a provider"""
    session_id = get_session_id(request)
    
    if provider_type not in ["gemini", "anthropic", "openai"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid provider type: {provider_type}. Must be 'gemini', 'anthropic', or 'openai'"
        )
    
    if not api_key or len(api_key.strip()) < 10:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid API key"
        )
    
    llm_manager.set_user_api_key(session_id, provider_type, api_key.strip())
    
    return JSONResponse(content={
        "success": True,
        "message": f"API key set for {provider_type}",
        "provider": provider_type
    })

@app.get("/api/get-api-keys")
async def get_api_keys(request: Request):
    """Get user API keys status (doesn't return actual keys for security)"""
    session_id = get_session_id(request)
    
    keys_status = {
        "gemini": session_id in llm_manager.user_api_keys and "gemini" in llm_manager.user_api_keys.get(session_id, {}),
        "anthropic": session_id in llm_manager.user_api_keys and "anthropic" in llm_manager.user_api_keys.get(session_id, {}),
        "openai": session_id in llm_manager.user_api_keys and "openai" in llm_manager.user_api_keys.get(session_id, {})
    }
    
    return JSONResponse(content={
        "has_user_keys": keys_status
    })

@app.post("/api/set-default-provider")
async def set_default_provider(
    request: Request,
    provider: str = Form(...),
    db: SQLASessionType = Depends(get_db)
):
    """Set default LLM provider/model"""
    session_id = get_session_id(request)
    
    if not llm_manager.is_model_available(provider, session_id):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Model '{provider}' not available. Please set an API key for this provider."
        )
    
    try:
        llm_manager.default_provider = LLMProvider(provider)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Invalid model: '{provider}'"
        )
    
    return JSONResponse(content={
        "success": True,
        "message": f"Default model set to {provider}",
        "provider": provider
    })

@app.post("/api/scrape-terminology")
async def scrape_terminology_endpoint(
    request: Request,
    admin_key: str = Form(...),
    db: SQLASessionType = Depends(get_db)
):
    """Admin endpoint to trigger scraping"""
    ADMIN_KEY = os.environ.get("ADMIN_KEY", "your-secret-admin-key-here")
    if admin_key != ADMIN_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    try:
        scraped_count = await scrape_eleto_documents()
        return JSONResponse(content={
            "success": True,
            "documents_scraped": scraped_count,
            "message": f"Successfully scraped {scraped_count} documents"
        })
    except Exception as e:
        logger.error(f"Scraping endpoint failed: {e}")
        raise HTTPException(status_code=500, detail=f"Scraping failed: {str(e)}")

@app.get("/api/find-exact/{term}")
async def find_exact_term(term: str):
    """Find exact occurrences of a term in all chunks"""
    matches = []
    term_lower = term.lower()
    
    for i, chunk in enumerate(rag_system.document_chunks):
        chunk_lower = chunk.lower()
        if term_lower in chunk_lower:
            # Find the context around the match
            start_pos = chunk_lower.find(term_lower)
            context_start = max(0, start_pos - 100)
            context_end = min(len(chunk), start_pos + len(term) + 100)
            context = chunk[context_start:context_end]
            
            matches.append({
                "chunk_index": i,
                "context": context,
                "full_chunk": chunk
            })
    
    return {
        "term": term,
        "found": len(matches) > 0,
        "matches": matches
    }

@app.get("/api/debug-chunks-vs-db")
async def debug_chunks_vs_db():
    """Compare in-memory chunks vs database content"""
    import sqlite3
    
    # Check database
    conn = sqlite3.connect('terminology_rag.db')
    cursor = conn.cursor()
    
    # Find chunks containing "taxiway" in database
    cursor.execute("SELECT chunk_text FROM documents WHERE chunk_text LIKE '%taxiway%' LIMIT 5")
    db_matches = [row[0] for row in cursor.fetchall()]
    
    # Find chunks containing "taxiway" in memory
    memory_matches = [chunk for chunk in rag_system.document_chunks if 'taxiway' in chunk.lower()]
    
    conn.close()
    
    return {
        "database_taxiway_matches": len(db_matches),
        "memory_taxiway_matches": len(memory_matches),
        "db_samples": db_matches[:2],
        "memory_samples": memory_matches[:2] if memory_matches else [],
        "total_chunks_in_memory": len(rag_system.document_chunks)
    }

@app.get("/api/direct-db-search/{term}")
async def direct_db_search(term: str):
    """Search directly in the database for debugging"""
    import sqlite3
    
    conn = sqlite3.connect('terminology_rag.db')
    cursor = conn.cursor()
    
    # Try multiple search patterns
    search_patterns = [
        f'%{term}%',                    # Original
        f'%{term.replace(" ", "%")}%',  # Words can be separated
        f'%{term.lower()}%',            # Lowercase
        f'%{term.upper()}%',            # Uppercase
    ]
    
    all_results = []
    
    for pattern in search_patterns:
        cursor.execute("SELECT chunk_text FROM documents WHERE LOWER(chunk_text) LIKE LOWER(?) LIMIT 5", (pattern,))
        results = cursor.fetchall()
        
        for result in results:
            chunk_text = result[0]
            if chunk_text not in [r["text"] for r in all_results]:  # Avoid duplicates
                all_results.append({
                    "text": chunk_text,
                    "pattern_matched": pattern,
                    "preview": chunk_text[:300]
                })
    
    conn.close()
    
    return {
        "term": term,
        "patterns_tried": search_patterns,
        "database_results": all_results,
        "count": len(all_results)
    }

@app.get("/api/test-line-chunking")
async def test_line_chunking():
    """Test line-based chunking on a sample"""
    
    sample_text = """χώρος στάθμευσης 
  
apron management service 
  
υπηρεσία διαχείρισης χώρων στάθμευσης 
  
apron taxiway 
  
τροχόδρομος χώρου στάθμευσης 
  
aquaplaning 
  
υδρολίσθηση

runway taxiway
  
διάδρομος τροχοδρόμου

taxiway intersection

διασταύρωση τροχοδρόμων"""
    
    # Test the new chunking method
    chunks = rag_system.chunk_text_by_lines(sample_text, min_lines=3, max_lines=8, overlap_lines=1)
    
    result = {
        "original_text": sample_text,
        "total_chunks": len(chunks),
        "chunks": []
    }
    
    for i, chunk in enumerate(chunks):
        result["chunks"].append({
            "chunk_id": i,
            "lines_count": len(chunk.split('\n')),
            "contains_apron_taxiway": "apron taxiway" in chunk.lower(),
            "chunk_text": chunk
        })
    
    return result

# Add this endpoint to rebuild with new chunking
@app.post("/api/rebuild-chunks")
async def rebuild_chunks():
    """Rebuild database with line-based chunking"""
    try:
        # Clear existing chunks
        conn = sqlite3.connect('terminology_rag.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM documents")
        conn.commit()
        conn.close()
        
        # Re-scrape and re-chunk
        scraped_count = await scrape_eleto_documents()
        return {"success": True, "documents_processed": scraped_count}
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.get("/api/rag-statistics")
async def get_rag_statistics():
    """Get statistics about RAG system sources"""
    if not rag_system.document_chunks:
        return {"total_chunks": 0, "sources": []}
    
    source_stats = {}
    for i, chunk in enumerate(rag_system.document_chunks):
        source = rag_system.chunk_sources[i] if i < len(rag_system.chunk_sources) else "Unknown"
        if source not in source_stats:
            source_stats[source] = {'name': source, 'chunk_count': 0}
        source_stats[source]['chunk_count'] += 1
    
    sources = []
    for source, stats in source_stats.items():
        doc_name = source.split('/')[-1] if '/' in source else source
        sources.append({
            'name': doc_name,
            'url': source,
            'chunk_count': stats['chunk_count']
        })
    
    return {
        "total_chunks": len(rag_system.document_chunks),
        "total_sources": len(sources),
        "sources": sorted(sources, key=lambda x: x['chunk_count'], reverse=True)
    }

@app.get("/api/debug-raw-db-content")
async def debug_raw_db_content():
    """Show raw database content for debugging"""
    import sqlite3
    
    conn = sqlite3.connect('terminology_rag.db')
    cursor = conn.cursor()
    
    # Get all chunks that contain "apron" OR "taxiway"
    cursor.execute("SELECT id, chunk_text FROM documents WHERE chunk_text LIKE '%apron%' OR chunk_text LIKE '%taxiway%' LIMIT 10")
    results = cursor.fetchall()
    
    debug_info = []
    for row_id, chunk_text in results:
        debug_info.append({
            "id": row_id,
            "length": len(chunk_text),
            "contains_apron": "apron" in chunk_text.lower(),
            "contains_taxiway": "taxiway" in chunk_text.lower(),
            "contains_both": "apron" in chunk_text.lower() and "taxiway" in chunk_text.lower(),
            "preview": chunk_text[:500],
            "raw_bytes": chunk_text.encode('utf-8', errors='replace')[:200]
        })
    
    conn.close()
    
    return {
        "total_results": len(results),
        "debug_info": debug_info
    }
@app.get("/api/search-terminology/{term}")
async def search_terminology_endpoint(term: str, k: int = 5):
    """Search for terminology using RAG system"""
    
    # Diagnostic logging
    logger.info(f"Search request for term: '{term}', k={k}")
    logger.info(f"RAG system status: chunks={len(rag_system.document_chunks) if rag_system and rag_system.document_chunks else 0}, index={'loaded' if rag_system and rag_system.index else 'not loaded'}")
    
    if not rag_system or not rag_system.document_chunks:
        error_msg = "RAG system not initialized or no documents loaded"
        logger.error(error_msg)
        return {
            "term": term,
            "total_results": 0,
            "results": [],
            "error": error_msg
        }
    
    try:
        # Use the RAG system's search function
        search_results = rag_system.search_relevant_content(term, k=k)
        logger.info(f"Search for '{term}' returned {len(search_results)} results")
        
        # Format results for frontend
        formatted_results = []
        for i, result in enumerate(search_results, 1):
            # Extract source filename from URL
            source_url = result.get('source', 'Unknown')
            source_file = source_url.split('/')[-1] if '/' in source_url else source_url
            
            formatted_results.append({
                'rank': i,
                'score': result.get('score', 0.0),
                'text': result.get('text', ''),
                'source': {
                    'file': source_file,
                    'url': source_url
                }
            })
        
        return {
            "term": term,
            "total_results": len(formatted_results),
            "results": formatted_results
        }
        
    except Exception as e:
        logger.error(f"Search failed for term '{term}': {e}", exc_info=True)
        return {
            "term": term,
            "total_results": 0,
            "results": [],
            "error": str(e)
        }
@app.get("/api/debug-embedding-search/{term}")
async def debug_embedding_search(term: str):
    """
    Diagnostic to perform a search directly with a known chunk's embedding.
    This bypasses the query embedding step to test the index integrity.
    """
    import sqlite3
    import pickle
    
    conn = sqlite3.connect('terminology_rag.db')
    cursor = conn.cursor()
    
    # Find a chunk that contains the term and has a stored embedding
    cursor.execute("SELECT chunk_text, embedding FROM documents WHERE chunk_text LIKE ? AND embedding IS NOT NULL LIMIT 1", (f'%{term}%',))
    result = cursor.fetchone()
    conn.close()
    
    if not result:
        raise HTTPException(status_code=404, detail=f"No chunk containing '{term}' with a valid embedding was found in the database.")
        
    source_chunk_text = result[0]
    source_embedding_bytes = result[1]
    
    try:
        source_embedding = pickle.loads(source_embedding_bytes)
        source_embedding = source_embedding.astype('float32').reshape(1, -1)
        faiss.normalize_L2(source_embedding)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load or process embedding from database: {str(e)}")

    if not rag_system.index:
        raise HTTPException(status_code=503, detail="FAISS index not loaded.")
    
    # Perform a search using this known, good embedding
    D, I = rag_system.index.search(source_embedding, k=5)
    
    found_results = []
    for i in range(len(I[0])):
        doc_idx = I[0][i]
        score = float(D[0][i])
        chunk_text = rag_system.document_chunks[doc_idx]
        found_results.append({"score": score, "text": chunk_text})

    return {
        "term": term,
        "source_chunk": source_chunk_text[:100] + "...",
        "faiss_results": found_results
    }

@app.post("/api/translate-with-terminology")
async def translate_with_terminology(
    request: Request,
    text: str = Form(...),
    source_language: str = Form("auto"),
    target_language: str = Form("en"),
    style: str = Form("formal"),
    llm_provider: str = Form(None),
    db: SQLASessionType = Depends(get_db)
):
    """Translate with terminology awareness using RAG"""
    current_user = get_current_user(request, db)
    
    if not terminology_translator:
        raise HTTPException(status_code=503, detail="Terminology translation service not available")
    
    try:
        if source_language == "auto":
            source_language = detect_text_language(text)
            
        provider = llm_provider or llm_manager.default_provider.value
        
        session_id = get_session_id(request)
        result = await terminology_translator.translate_with_rag(
            text, source_language, target_language, style, provider, db, session_id
        )
        
        # FIXED: Return the full result object, not nested
        return JSONResponse(content={
            "success": True,
            "translation": result['translated_text'],
            "source_language": result['source_language'],
            "target_language": result['target_language'],
            "style": style,
            "llm_provider": result.get('llm_provider', provider),
            "confidence_score": result.get('confidence_score', 0.5),
            "rag_sources_used": result.get('sources_used', 0),
            "sources": result.get('sources', [])
        })
        
    except Exception as e:
        logger.error(f"Terminology translation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")
@app.get("/api/sample-chunks")
async def sample_chunks():
    return {
        "sample_chunks": [
            chunk[:2000] for chunk in rag_system.document_chunks[:5]
        ]
    }
@app.get("/", response_class=HTMLResponse)
async def serve_index():
    # --- FIX: Use absolute path for templates ---
    index_path = TEMPLATES_DIR / "index.html"
    with open(index_path, "r", encoding="utf-8") as f:
        return HTMLResponse(f.read())
@app.post("/api/terminology-chat")
async def terminology_chat_endpoint(
    request: Request,
    message: str = Form(...),
    conversation_history: str = Form("[]"),
    llm_provider: str = Form(None),
    db: SQLASessionType = Depends(get_db)
):
    """Chat about terminology with RAG enhancement"""
    current_user = get_current_user(request, db)
    
    if not terminology_chat:
        raise HTTPException(status_code=503, detail="Terminology chat service not available")
    
    try:
        history = json.loads(conversation_history) if conversation_history else []
        provider = llm_provider or llm_manager.default_provider.value
        session_id = get_session_id(request)
        result = await terminology_chat.chat_about_terminology(
            user_message=message, 
            conversation_history=history, 
            provider=provider, 
            db=db,
            session_id=session_id
        )      

        return JSONResponse(content={
            "success": True,
            "response": result['response'],
            "sources": result.get('sources', []),
            "llm_provider": provider
        })
        
    except Exception as e:
        logger.error(f"Terminology chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")

@app.post("/analyze-text-multi")
async def analyze_text_multi(
    request: Request,
    text: str = Form(...), 
    style_type: str = Form("academic"), 
    language: str = Form("en"),
    llm_provider: str = Form(None),
    db: SQLASessionType = Depends(get_db)
):
    """Analyze text with LLM provider selection"""
    current_user = get_current_user(request, db)
    
    try:
        if not text or len(text.strip()) < 10:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST, 
                detail="Please provide at least 10 characters of text"
            )
        
        text = text.strip()
        provider = llm_provider or llm_manager.default_provider.value
        
        logger.info(f"Analyzing {len(text)} characters using {provider}")
        
        if language == "el":
            prompt = f"""Είστε ένας ειδικός αναλυτής γραφής. Αναλύστε αυτό το κείμενο για το στυλ γραφής {style_type}.

Κείμενο: "{text}"

Παρέχετε ανάλυση σε μορφή JSON:
{{
    "grammar_issues": [
        {{"issue": "περιγραφή προβλήματος", "suggestion": "σύσταση διόρθωσης", "severity": "high/medium/low"}}
    ],
    "style_analysis": {{
        "tone": "τόνος και χαρακτηριστικά στυλ",
        "clarity_score": "1-10",
        "style_alignment": "ευθυγράμμιση με το στυλ {style_type}"
    }},
    "suggestions": ["συστάσεις βελτίωσης"],
    "llm_used": "{provider}"
}}"""
        else:
            prompt = f"""You are an expert writing analyst. Analyze this text for {style_type} writing style.

Text: "{text}"

Provide analysis in JSON format:
{{
    "grammar_issues": [
        {{"issue": "problem description", "suggestion": "correction recommendation", "severity": "high/medium/low"}}
    ],
    "style_analysis": {{
        "tone": "detected tone and style characteristics",
        "clarity_score": "1-10",
        "style_alignment": "assessment of alignment with {style_type} style"
    }},
    "suggestions": ["improvement recommendations"],
    "llm_used": "{provider}"
}}"""
        
        session_id = get_session_id(request)
        result = await generate_with_timeout_multi(prompt, provider, timeout=45, session_id=session_id)
        
        try:
            parsed_result = json.loads(result)
            parsed_result["llm_provider"] = provider
            return parsed_result
        except json.JSONDecodeError:
            logger.warning(f"AI response was not valid JSON. Raw response: {result[:200]}...")
            return {
                "raw_analysis": result, 
                "status": "success",
                "llm_provider": provider
            }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
@app.post("/api/favorites/add")
async def add_favorite(
    translation_id: int = Form(...),
    notes: str = Form(""),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Add translation to favorites"""
    try:
        session = get_extended_session()
        
        # Update translation as favorite
        translation = session.query(TranslationHistory).filter_by(id=translation_id).first()
        if translation:
            translation.is_favorite = True
            
        # Create favorite entry
        favorite = FavoriteTranslation(
            translation_id=translation_id,
            notes=notes,
            created_at=datetime.utcnow()
        )
        session.add(favorite)
        session.commit()
        
        return {"success": True, "message": "Added to favorites"}
    except Exception as e:
        logger.error(f"Add favorite failed: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/favorites/list")
async def list_favorites(
    limit: int = 50,
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """List all favorite translations"""
    try:
        session = get_extended_session()
        favorites = session.query(TranslationHistory).filter_by(is_favorite=True).order_by(TranslationHistory.created_at.desc()).limit(limit).all()
        
        result = []
        for fav in favorites:
            result.append({
                'id': fav.id,
                'source_text': fav.source_text,
                'translated_text': fav.translated_text,
                'source_language': fav.source_language,
                'target_language': fav.target_language,
                'created_at': fav.created_at.isoformat()
            })
        
        return {"success": True, "favorites": result}
    except Exception as e:
        logger.error(f"List favorites failed: {e}")
        return {"success": False, "error": str(e)}

# === FEATURE 2: Translation History ===

@app.post("/api/history/save")
async def save_translation_history(
    source_text: str = Form(...),
    translated_text: str = Form(...),
    source_language: str = Form(...),
    target_language: str = Form(...),
    style: str = Form("formal"),
    llm_provider: str = Form(""),
    confidence_score: float = Form(0.0),
    sources_used: int = Form(0),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Save translation to history"""
    try:
        session = get_extended_session()
        
        history = TranslationHistory(
            source_text=source_text,
            translated_text=translated_text,
            source_language=source_language,
            target_language=target_language,
            style=style,
            llm_provider=llm_provider,
            confidence_score=confidence_score,
            sources_used=sources_used,
            created_at=datetime.utcnow()
        )
        session.add(history)
        session.commit()
        
        return {"success": True, "id": history.id}
    except Exception as e:
        logger.error(f"Save history failed: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/history/list")
async def list_translation_history(
    limit: int = 100,
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """List translation history"""
    try:
        session = get_extended_session()
        history = session.query(TranslationHistory).order_by(TranslationHistory.created_at.desc()).limit(limit).all()
        
        result = []
        for h in history:
            result.append({
                'id': h.id,
                'source_text': h.source_text[:100] + '...' if len(h.source_text) > 100 else h.source_text,
                'translated_text': h.translated_text[:100] + '...' if len(h.translated_text) > 100 else h.translated_text,
                'source_language': h.source_language,
                'target_language': h.target_language,
                'llm_provider': h.llm_provider,
                'is_favorite': h.is_favorite,
                'created_at': h.created_at.isoformat()
            })
        
        return {"success": True, "history": result}
    except Exception as e:
        logger.error(f"List history failed: {e}")
        return {"success": False, "error": str(e)}

# === FEATURE 3: Custom Glossary ===
@app.post("/api/glossary/add")
async def add_glossary_term(
    source_term: str = Form(...),
    target_term: str = Form(...),
    source_language: str = Form(...),
    target_language: str = Form(...),
    context: str = Form(""),
    priority: int = Form(1),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Add term to custom glossary"""
    try:
        term = CustomGlossary(
            source_term=source_term,
            target_term=target_term,
            source_language=source_language,
            target_language=target_language,
            context=context,
            priority=priority,
            created_at=datetime.utcnow()
        )
        db.add(term)
        db.commit()
        db.refresh(term)
        term_id = term.id
        
        return {"success": True, "id": term_id}
    except Exception as e:
        db.rollback()
        logger.error(f"Add glossary term failed: {e}")
        return {"success": False, "error": str(e)}
@app.get("/api/glossary/search")
async def search_glossary(
    term: str,
    source_lang: str,
    target_lang: str,
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Search custom glossary for a term"""
    try:
        # Exact match first
        exact = db.query(CustomGlossary).filter(
            CustomGlossary.source_term.ilike(term),
            CustomGlossary.source_language == source_lang,
            CustomGlossary.target_language == target_lang
        ).order_by(CustomGlossary.priority.desc()).first()
        
        if exact:
            return {
                "success": True,
                "found": True,
                "source_term": exact.source_term,
                "target_term": exact.target_term,
                "context": exact.context
            }
        
        # Partial match
        partial = db.query(CustomGlossary).filter(
            CustomGlossary.source_term.ilike(f'%{term}%'),
            CustomGlossary.source_language == source_lang,
            CustomGlossary.target_language == target_lang
        ).order_by(CustomGlossary.priority.desc()).limit(5).all()
        
        results = [{
            "source_term": p.source_term,
            "target_term": p.target_term,
            "context": p.context
        } for p in partial]
        
        return {"success": True, "found": len(results) > 0, "matches": results}
    except Exception as e:
        logger.error(f"Search glossary failed: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/glossary/list")
async def list_glossary(
    source_lang: str = None,
    target_lang: str = None,
    limit: int = 100,
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """List all glossary terms"""
    try:
        query = db.query(CustomGlossary)
        
        if source_lang:
            query = query.filter_by(source_language=source_lang)
        if target_lang:
            query = query.filter_by(target_language=target_lang)
        
        terms = query.order_by(CustomGlossary.priority.desc(), CustomGlossary.source_term).limit(limit).all()
        
        result = [{
            'id': t.id,
            'source_term': t.source_term,
            'target_term': t.target_term,
            'source_language': t.source_language,
            'target_language': t.target_language,
            'context': t.context,
            'priority': t.priority
        } for t in terms]
        
        return {"success": True, "terms": result}
    except Exception as e:
        logger.error(f"List glossary failed: {e}")
        return {"success": False, "error": str(e)}

@app.delete("/api/glossary/delete/{term_id}")
async def delete_glossary_term(
    term_id: int,
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Delete glossary term"""
    try:
        term = db.query(CustomGlossary).filter_by(id=term_id).first()
        if term:
            db.delete(term)
            db.commit()
            return {"success": True}
        return {"success": False, "error": "Term not found"}
    except Exception as e:
        logger.error(f"Delete glossary term failed: {e}")
        return {"success": False, "error": str(e)}
# === FEATURE 4: Export to PDF/DOCX ===

@app.post("/api/export/translation")
async def export_translation(
    translation_id: int = Form(...),
    format: str = Form("docx"),  # docx or pdf
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Export translation to DOCX or PDF"""
    try:
        session = get_extended_session()
        translation = session.query(TranslationHistory).filter_by(id=translation_id).first()
        
        if not translation:
            raise HTTPException(status_code=404, detail="Translation not found")
        
        if format == "docx":
            # Create DOCX
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Translation', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Date: {translation.created_at.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"Languages: {translation.source_language} → {translation.target_language}")
            doc.add_paragraph(f"Provider: {translation.llm_provider}")
            doc.add_paragraph(f"Confidence: {translation.confidence_score * 100:.1f}%")
            
            doc.add_paragraph()  # Spacing
            
            # Source text
            doc.add_heading('Source Text', 2)
            p = doc.add_paragraph(translation.source_text)
            p.style = 'Normal'
            
            doc.add_paragraph()  # Spacing
            
            # Translation
            doc.add_heading('Translation', 2)
            p = doc.add_paragraph(translation.translated_text)
            p.style = 'Normal'
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=translation_{translation_id}.docx"}
            )
        
        elif format == "pdf":
            # For PDF, we'll create a simple HTML and convert
            # (You'll need to install reportlab or weasyprint)
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            output = BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            story.append(Paragraph("Translation", styles['Title']))
            story.append(Spacer(1, 12))
            
            # Metadata
            story.append(Paragraph(f"Date: {translation.created_at.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
            story.append(Paragraph(f"Languages: {translation.source_language} → {translation.target_language}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Source
            story.append(Paragraph("Source Text:", styles['Heading2']))
            story.append(Paragraph(translation.source_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Translation
            story.append(Paragraph("Translation:", styles['Heading2']))
            story.append(Paragraph(translation.translated_text, styles['Normal']))
            
            doc.build(story)
            output.seek(0)
            
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=translation_{translation_id}.pdf"}
            )
        
        else:
            raise HTTPException(status_code=400, detail="Invalid format")
            
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === FEATURE 5: Batch Translation ===

@app.post("/api/batch/upload")
async def batch_upload(
    file: UploadFile = File(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    style: str = Form("formal"),
    llm_provider: str = Form(None),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Upload file for batch translation (CSV, TXT)"""
    try:
        # Read file
        content = await file.read()
        content_str = content.decode('utf-8')
        
        # Parse based on file type
        lines = []
        if file.filename.endswith('.csv'):
            csv_reader = csv.reader(StringIO(content_str))
            lines = [row[0] for row in csv_reader if row]
        else:
            lines = content_str.split('\n')
        
        lines = [line.strip() for line in lines if line.strip()]
        
        # Create batch record
        session = get_extended_session()
        batch = BatchTranslation(
            filename=file.filename,
            source_language=source_lang,
            target_language=target_lang,
            total_items=len(lines),
            status='pending'
        )
        session.add(batch)
        session.commit()
        
        # Start async translation
        asyncio.create_task(process_batch_translation(batch.id, lines, source_lang, target_lang, style, llm_provider))
        
        return {
            "success": True,
            "batch_id": batch.id,
            "total_items": len(lines),
            "message": "Batch translation started"
        }
        
    except Exception as e:
        logger.error(f"Batch upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def process_batch_translation(batch_id, lines, source_lang, target_lang, style, provider):
    """Process batch translation in background"""
    session = get_extended_session()
    batch = session.query(BatchTranslation).filter_by(id=batch_id).first()
    batch.status = 'processing'
    session.commit()
    
    results = []
    provider = provider or llm_manager.default_provider.value
    
    for i, line in enumerate(lines):
        try:
            result = await terminology_translator.translate_with_rag(
                line, source_lang, target_lang, style, provider
            )
            results.append({
                'source': line,
                'translation': result['translated_text']
            })
            
            batch.completed_items = i + 1
            session.commit()
            
        except Exception as e:
            logger.error(f"Batch item {i} failed: {e}")
            results.append({
                'source': line,
                'translation': f"[ERROR: {str(e)}]"
            })
    
    # Save results as CSV
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(['Source', 'Translation'])
    for r in results:
        writer.writerow([r['source'], r['translation']])
    
    # Save to file
    output_path = f"batch_output_{batch_id}.csv"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(output.getvalue())
    
    batch.status = 'completed'
    batch.output_file_path = output_path
    batch.completed_at = datetime.utcnow()
    session.commit()

@app.get("/api/batch/status/{batch_id}")
async def batch_status(batch_id: int):
    """Check batch translation status"""
    try:
        session = get_extended_session()
        batch = session.query(BatchTranslation).filter_by(id=batch_id).first()
        
        if not batch:
            raise HTTPException(status_code=404, detail="Batch not found")
        
        return {
            "success": True,
            "status": batch.status,
            "completed": batch.completed_items,
            "total": batch.total_items,
            "progress": (batch.completed_items / batch.total_items * 100) if batch.total_items > 0 else 0
        }
    except Exception as e:
        logger.error(f"Batch status failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/batch/download/{batch_id}")
async def batch_download(batch_id: int):
    """Download completed batch translation"""
    try:
        session = get_extended_session()
        batch = session.query(BatchTranslation).filter_by(id=batch_id).first()
        
        if not batch or batch.status != 'completed':
            raise HTTPException(status_code=404, detail="Batch not ready")
        
        return FileResponse(
            batch.output_file_path,
            media_type="text/csv",
            filename=f"translation_{batch_id}.csv"
        )
    except Exception as e:
        logger.error(f"Batch download failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === JSON FILE UPLOAD FEATURE ===

@app.post("/api/upload-json")
async def upload_json_file(
    file: UploadFile = File(...),
    upload_type: str = Form(...),
    json_data: str = Form(None),
    db: SQLASessionType = Depends(get_db)
):
    """
    Upload and process JSON files for:
    - Glossary import
    - Batch translation
    - Custom terminology
    """
    try:
        # Read and parse JSON file
        content = await file.read()
        
        try:
            if json_data:
                data = json.loads(json_data)
            else:
                content_str = content.decode('utf-8')
                data = json.loads(content_str)
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON: {e}")
            return JSONResponse({
                "success": False,
                "detail": f"Invalid JSON format: {str(e)}"
            }, status_code=400)
        
        # Auto-detect upload type if set to auto
        if upload_type == "auto":
            upload_type = detect_json_type(data)
            logger.info(f"Auto-detected JSON type: {upload_type}")
        
        # Process based on type
        if upload_type == "glossary":
            result = await process_glossary_json(data, db)
        elif upload_type == "batch_translate":
            result = await process_batch_translate_json(data, db)
        elif upload_type == "terminology":
            result = await process_terminology_json(data, db)
        else:
            return JSONResponse({
                "success": False,
                "detail": f"Unknown upload type: {upload_type}"
            }, status_code=400)
        
        return JSONResponse({
            "success": True,
            "type": upload_type,
            "message": result.get("message", "JSON processed successfully"),
            **result
        })
        
    except Exception as e:
        logger.error(f"JSON upload failed: {e}")
        return JSONResponse({
            "success": False,
            "detail": str(e)
        }, status_code=500)

def detect_json_type(data: dict) -> str:
    """Auto-detect JSON file type based on structure"""
    if "glossary" in data or ("source" in str(data) and "target" in str(data)):
        return "glossary"
    elif "items" in data and isinstance(data.get("items"), list):
        return "batch_translate"
    elif "terms" in data or ("term_greek" in str(data) and "term_english" in str(data)):
        return "terminology"
    else:
        # Default to glossary
        return "glossary"

async def process_glossary_json(data: dict, db: SQLASessionType) -> dict:
    """Process glossary import JSON - flexible format support"""
    imported = 0
    skipped = 0
    errors = 0
    sample = []
    
    try:
        # Support multiple formats
        glossary_items = []
        
        if "glossary" in data:
            glossary_items = data["glossary"]
        elif isinstance(data, list):
            glossary_items = data
        else:
            glossary_items = [data]
        
        for item in glossary_items:
            try:
                # Flexible field name mapping - support various naming conventions
                source_term = (
                    item.get("source_term") or 
                    item.get("source") or 
                    item.get("term") or
                    item.get("sourceterm")
                )
                
                target_term = (
                    item.get("target_term") or 
                    item.get("target") or 
                    item.get("translation") or
                    item.get("targetterm")
                )
                
                if not source_term or not target_term:
                    logger.warning(f"Missing required fields in item: {item}")
                    errors += 1
                    continue
                
                # Auto-detect language direction based on script
                # Greek uses Greek alphabet, English uses Latin alphabet
                def detect_language(text):
                    # Check if text contains Greek characters
                    has_greek = bool(re.search(r'[Α-Ωα-ω]', text))
                    return 'el' if has_greek else 'en'
                
                # Get or detect languages
                source_lang = (
                    item.get("source_lang") or 
                    item.get("source_language") or 
                    item.get("sourcelang") or
                    detect_language(source_term)
                )
                
                target_lang = (
                    item.get("target_lang") or 
                    item.get("target_language") or 
                    item.get("targetlang") or
                    detect_language(target_term)
                )
                
                # Get context/domain information
                context = (
                    item.get("context") or 
                    item.get("description") or 
                    item.get("note") or 
                    ""
                )
                
                domain = item.get("domain") or item.get("category") or ""
                
                # Combine context and domain if both exist
                if domain and context:
                    full_context = f"{context} | Domain: {domain}"
                elif domain:
                    full_context = f"Domain: {domain}"
                else:
                    full_context = context
                
                # Check if already exists
                existing = db.query(CustomGlossary).filter_by(
                    source_term=source_term,
                    target_term=target_term,
                    source_language=source_lang,
                    target_language=target_lang
                ).first()
                
                if existing:
                    skipped += 1
                    logger.info(f"Skipped duplicate: {source_term} -> {target_term}")
                    continue
                
                # Add to glossary
                new_entry = CustomGlossary(
                    source_term=source_term,
                    target_term=target_term,
                    source_language=source_lang,
                    target_language=target_lang,
                    context=full_context,
                    priority=item.get("priority", 1),
                    created_at=datetime.utcnow()
                )
                
                db.add(new_entry)
                imported += 1
                
                if len(sample) < 5:
                    sample.append({
                        "source": source_term,
                        "target": target_term,
                        "lang": f"{source_lang} → {target_lang}",
                        "domain": domain or "general"
                    })
                
            except Exception as e:
                logger.error(f"Error processing glossary item: {e}")
                errors += 1
                continue
        
        db.commit()
        logger.info(f"Glossary import complete: {imported} imported, {skipped} skipped, {errors} errors")
        
        # ALSO ADD TO RAG INDEX for enhanced semantic search
        added_to_rag = False
        try:
            if imported > 0:
                # Create RAG-friendly documents from glossary
                rag_documents = []
                rag_sources = []
                
                for item in glossary_items[:imported]:  # Only process successfully imported items
                    source = item.get("source_term") or item.get("source")
                    target = item.get("target_term") or item.get("target")
                    context = item.get("context", "")
                    domain = item.get("domain", "custom")
                    
                    if source and target:
                        # Create searchable text
                        doc_text = f"{source} | {target}"
                        if context:
                            doc_text += f"\n{context}"
                        if domain:
                            doc_text += f"\nDomain: {domain}"
                        
                        rag_documents.append(doc_text)
                        rag_sources.append({
                            "file": f"custom_glossary_{domain}",
                            "url": "user_uploaded_glossary",
                            "domain": domain
                        })
                
                # Add to RAG system
                if rag_documents and hasattr(rag_system, 'add_documents'):
                    rag_system.add_documents(rag_documents, rag_sources)
                    added_to_rag = True
                    logger.info(f"Added {len(rag_documents)} glossary terms to RAG index")
                
        except Exception as e:
            logger.warning(f"Failed to add glossary to RAG (terms still in database): {e}")
        
        return {
            "imported": imported,
            "skipped": skipped,
            "errors": errors,
            "sample": sample,
            "added_to_rag": added_to_rag,
            "message": f"Successfully imported {imported} glossary terms ({skipped} skipped, {errors} errors)" + 
                      (f" + Added to RAG index" if added_to_rag else "")
        }
        
    except Exception as e:
        logger.error(f"Glossary JSON processing failed: {e}")
        db.rollback()
        raise

async def process_batch_translate_json(data: dict, db: SQLASessionType) -> dict:
    """Process batch translation JSON"""
    try:
        items = data.get("items", [])
        settings = data.get("settings", {})
        
        if not items:
            raise ValueError("No items found in JSON")
        
        # Extract settings
        source_lang = settings.get("source_lang", "el")
        target_lang = settings.get("target_lang", "en")
        style = settings.get("style", "formal")
        llm_provider = settings.get("llm_provider", llm_manager.default_provider.value)
        
        # Create batch translation entry
        batch = BatchTranslation(
            user_id=None,
            source_language=source_lang,
            target_language=target_lang,
            style=style,
            llm_provider=llm_provider,
            total_items=len(items),
            status='pending',
            created_at=datetime.utcnow()
        )
        
        db.add(batch)
        db.commit()
        db.refresh(batch)
        
        # Store items for processing
        batch_items = []
        for idx, item in enumerate(items):
            text = item.get("text", "")
            item_source_lang = item.get("source_lang", source_lang)
            item_target_lang = item.get("target_lang", target_lang)
            
            if text:
                batch_items.append({
                    "index": idx,
                    "text": text,
                    "source_lang": item_source_lang,
                    "target_lang": item_target_lang
                })
        
        # Start async processing (you would implement this)
        # For now, just return batch info
        
        return {
            "batch_id": batch.id,
            "total_items": len(items),
            "message": f"Batch translation started with {len(items)} items"
        }
        
    except Exception as e:
        logger.error(f"Batch translate JSON processing failed: {e}")
        db.rollback()
        raise

async def process_terminology_json(data: dict, db: SQLASessionType) -> dict:
    """Process custom terminology JSON and add to RAG system"""
    imported = 0
    sample = []
    
    try:
        # Support multiple formats
        term_items = []
        
        if "terms" in data:
            term_items = data["terms"]
        elif isinstance(data, list):
            term_items = data
        else:
            term_items = [data]
        
        # Process and add to RAG
        documents = []
        sources = []
        
        for item in term_items:
            try:
                term_greek = item.get("term_greek", "")
                term_english = item.get("term_english", "")
                definition = item.get("definition", "") or item.get("definition_greek", "")
                domain = item.get("domain", "custom")
                
                if not term_greek or not term_english:
                    continue
                
                # Create document text for RAG
                doc_text = f"{term_greek} | {term_english}\n{definition}"
                documents.append(doc_text)
                
                sources.append({
                    "file": f"custom_terminology_{domain}",
                    "url": "custom_import",
                    "domain": domain
                })
                
                imported += 1
                
                if len(sample) < 5:
                    sample.append({
                        "greek": term_greek,
                        "english": term_english,
                        "domain": domain
                    })
                
            except Exception as e:
                logger.error(f"Error processing terminology item: {e}")
                continue
        
        # Add to RAG system
        added_to_rag = False
        if documents and hasattr(rag_system, 'add_documents'):
            try:
                rag_system.add_documents(documents, sources)
                added_to_rag = True
            except Exception as e:
                logger.error(f"Failed to add to RAG: {e}")
        
        return {
            "imported": imported,
            "added_to_rag": added_to_rag,
            "sample": sample,
            "message": f"Imported {imported} terminology entries"
        }
        
    except Exception as e:
        logger.error(f"Terminology JSON processing failed: {e}")
        raise

# === FEATURE 6: Term Extraction ===

@app.post("/api/extract/terms")
async def extract_terms(
    file: UploadFile = File(None),
    text: str = Form(None),
    language: str = Form("el"),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Extract terms from text using LLM"""
    try:
        # Get text
        if file:
            content = await file.read()
            if file.filename.endswith('.pdf'):
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            elif file.filename.endswith('.docx'):
                import docx
                doc = docx.Document(BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                text = content.decode('utf-8')
        
        if not text:
            raise HTTPException(status_code=400, detail="No text provided")
        
        # Use LLM to extract technical terms
        prompt = f"""Extract the 20 most important technical terms from this Greek text. Return ONLY a JSON array of terms, nothing else.

Text:
{text[:2000]}

Example output: ["τεχνητή νοημοσύνη", "μηχανική μάθηση", "αλγόριθμος"]

JSON array:"""

        session_id = get_session_id(request) if request else None
        try:
            llm_response = await generate_with_timeout_multi(prompt, llm_manager.default_provider.value, timeout=30, session_id=session_id)
            
            # Parse JSON response
            import re
            json_match = re.search(r'\[.*\]', llm_response, re.DOTALL)
            if json_match:
                terms_list = json.loads(json_match.group())
            else:
                terms_list = []
        except:
            terms_list = []
        
        # For each term, search RAG for translation
        found_terms = []
        
        for term in terms_list[:20]:
            results = rag_system.search_relevant_content(term, k=5)
            
            if results and len(results) > 0:
                # Get top 5 chunks
                chunks = [r.get('text', '') for r in results[:5]]
                combined = "\n\n".join(chunks)
                
                # Ask LLM to find translation
                trans_prompt = f"""Find the English translation for the Greek term "{term}" in these text chunks.

Chunks:
{combined[:1500]}

If you find a translation, respond with ONLY the English term, nothing else. If no translation found, respond with "NONE"."""

                try:
                    translation = await generate_with_timeout_multi(trans_prompt, llm_manager.default_provider.value, timeout=20, session_id=session_id)
                    translation = translation.strip()
                    
                    if translation and translation != "NONE" and len(translation) < 100:
                        source_url = results[0].get('source', '')
                        found_terms.append({
                            'term': term,
                            'translation': translation,
                            'score': round(results[0].get('score', 0) * 100, 1),
                            'source': source_url.split('/')[-1] if '/' in source_url else source_url,
                            'source_url': source_url
                        })
                except:
                    continue
        
        return {
            "success": True,
            "terms": found_terms,
            "total_found": len(found_terms)
        }
        
    except Exception as e:
        logger.error(f"Term extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
@app.get("/api/suggestions/{text}")
async def get_suggestions(
    text: str,
    source_lang: str = "el",
    target_lang: str = "en",
    limit: int = 5
):
    """Get terminology suggestions as user types"""
    try:
        if len(text) < 3:
            return {"success": True, "suggestions": []}
        
        # Check custom glossary first
        session = get_extended_session()
        glossary_matches = session.query(CustomGlossary).filter(
            CustomGlossary.source_term.ilike(f'%{text}%'),
            CustomGlossary.source_language == source_lang,
            CustomGlossary.target_language == target_lang
        ).limit(3).all()
        
        suggestions = [{
            "source": g.source_term,
            "target": g.target_term,
            "type": "glossary"
        } for g in glossary_matches]
        
        # Then check RAG system
        if len(suggestions) < limit:
            chunks = rag_system.search_relevant_content(text, k=limit - len(suggestions))
            for chunk_data in chunks:
                chunk_text = chunk_data.get('text', '')
                # Extract relevant snippet
                if text.lower() in chunk_text.lower():
                    start = max(0, chunk_text.lower().find(text.lower()) - 50)
                    end = min(len(chunk_text), start + 150)
                    snippet = chunk_text[start:end]
                    suggestions.append({
                        "snippet": snippet,
                        "source_file": chunk_data.get('source', ''),
                        "type": "rag"
                    })
        
        return {"success": True, "suggestions": suggestions[:limit]}
        
    except Exception as e:
        logger.error(f"Suggestions failed: {e}")
        return {"success": False, "error": str(e)}

# === FEATURE 8: Translation Quality Scoring ===

@app.post("/api/quality/score")
async def score_translation_quality(
    translations: str = Form(...),  # JSON array of translation IDs
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Score consistency of terminology across translations"""
    try:
        import json
        translation_ids = json.loads(translations)
        
        session = get_extended_session()
        trans = session.query(TranslationHistory).filter(
            TranslationHistory.id.in_(translation_ids)
        ).all()
        
        if len(trans) < 2:
            return {"success": False, "error": "Need at least 2 translations"}
        
        # Analyze term consistency
        all_terms = {}
        for t in trans:
            words = set(t.translated_text.lower().split())
            for word in words:
                if len(word) > 4:  # Only significant words
                    all_terms[word] = all_terms.get(word, 0) + 1
        
        # Find consistent terms (appear in multiple translations)
        consistent_terms = {k: v for k, v in all_terms.items() if v > 1}
        consistency_score = len(consistent_terms) / len(all_terms) if all_terms else 0
        
        # Check confidence scores
        avg_confidence = sum(t.confidence_score or 0 for t in trans) / len(trans)
        
        # Overall quality score
        quality_score = (consistency_score * 0.6 + avg_confidence * 0.4)
        
        return {
            "success": True,
            "quality_score": quality_score * 100,
            "consistency_score": consistency_score * 100,
            "avg_confidence": avg_confidence * 100,
            "consistent_terms": len(consistent_terms),
            "total_unique_terms": len(all_terms)
        }
        
    except Exception as e:
        logger.error(f"Quality scoring failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === FEATURE 9: Multi-LLM Comparison ===

@app.post("/api/compare/translate")
async def compare_translation(
    text: str = Form(...),
    source_lang: str = Form("el"),
    target_lang: str = Form("en"),
    style: str = Form("formal"),
    request: Request = None,
    db: SQLASessionType = Depends(get_db)
):
    """Translate with all available LLMs and compare"""
    try:
        providers = llm_manager.get_available_providers()
        
        results = []
        for provider_key in providers.keys():
            try:
                result = await terminology_translator.translate_with_rag(
                    text, source_lang, target_lang, style, provider_key
                )
                results.append({
                    "provider": provider_key,
                    "provider_name": providers[provider_key]["name"],
                    "translation": result['translated_text'],
                    "confidence": result.get('confidence_score', 0),
                    "sources_used": result.get('sources_used', 0)
                })
            except Exception as e:
                logger.error(f"Provider {provider_key} failed: {e}")
                results.append({
                    "provider": provider_key,
                    "provider_name": providers[provider_key]["name"],
                    "translation": f"[Error: {str(e)}]",
                    "confidence": 0,
                    "sources_used": 0
                })
        
        return {
            "success": True,
            "source_text": text,
            "comparisons": results
        }
        
    except Exception as e:
        logger.error(f"Comparison failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
# Application startup
@app.on_event("startup")
async def startup_event():
    """Initialize services on startup"""
    create_tables()
    logger.info("SUCCESS: Database tables created")
    
    # Initialize terminology system
    initialize_terminology_system()
    logger.info("SUCCESS: Terminology system initialized")

@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup on shutdown"""
    await scraper.close_session()
    logger.info("SUCCESS: Scraper session closed")
@app.get("/api/debug-in-memory-chunks/{term}")
async def debug_in_memory_chunks(term: str):
    """
    Checks if a direct match exists in the in-memory document chunks list.
    """
    term_lower = term.lower().strip()
    matches = [chunk for chunk in rag_system.document_chunks if term_lower in chunk.lower()]
    
    return {
        "term": term,
        "found_in_memory": len(matches) > 0,
        "total_chunks_in_memory": len(rag_system.document_chunks),
        "matches": matches[:3]
    }
if __name__ == "__main__":
    rag_system.load_embedding_model()
    rag_system.build_faiss_index()
    print("FAISS index has been rebuilt from files. Check log for confirmation.")